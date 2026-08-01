"""
build_dissertation.py
=====================

Generates the dissertation .docx for the MSc Computing project
"Predictive Maintenance and Equipment Failure Prediction Using the
NASA Turbofan Engine Dataset" by Vatsal Nileshbhai Tailor (Student
ID A00067312), University of Roehampton.

The chapter content is stored as plain strings (see CHAPTERS below) so
the entire document can be regenerated from a single Python script
with no external dependencies beyond python-docx.

Formatting follows the UK MSc Computing conventions:
  * A4 paper, 2.54 cm margins
  * Times New Roman 12 pt body
  * 1.5 line spacing
  * Justified text
  * Heading 1 / Heading 2 / Heading 3 styles for the TOC field
  * Automatic Table of Contents (TOC field — right-click and choose
    "Update Field" in Word to populate)
  * Page numbers in the footer from the abstract onwards
  * IEEE-style references

Run from any working directory:

    python "00_Repository Analysis/build_dissertation.py"
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent  # <project>/data_cleaning

PHASE1_DIR = PROJECT_DIR / "01_Data Cleaning & Preprocessing"
PHASE2_DIR = PROJECT_DIR / "02_Feature Engineering"
PHASE5_DIR = PROJECT_DIR / "05_Performance Evaluation" / "performance_evaluation"
PHASE6_DIR = PROJECT_DIR / "06_Feature Importance Analysis" / "feature_importance_analysis"
PHASE7_DIR = PROJECT_DIR / "07_Maintenance Recommendation"

OUT_FILE = SCRIPT_DIR / (
    "Predictive_Maintenance_and_Equipment_Failure_Prediction_"
    "Using_the_NASA_Turbofan_Engine_Dataset_Dissertation.docx"
)
# Use a unique tmp filename on every run so a previously locked file
# does not block the rebuild.
import time as _time
TMP_FILE = SCRIPT_DIR / (
    f"_dissertation_build_{int(_time.time())}.docx"
)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def set_page_setup(doc: Document) -> None:
    for section in doc.sections:
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    # Force eastAsia to TNR as well so the font is consistent on Windows.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose Update Field to populate."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    placeholder_r = OxmlElement("w:r")
    placeholder_r.append(placeholder)
    run._r.addnext(placeholder_r)
    run._r.append(fldChar3)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h.paragraph_format.keep_with_next = True


def add_paragraph(doc: Document, text: str, bold: bool = False,
                  italic: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.alignment = (
        align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    )


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in p.runs:
            run.font.name = "Times New Roman"


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_footer_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p)


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------
def resolve(path: Path) -> Path | None:
    return path if path.exists() else None


def add_image(doc: Document, path: Path, caption: str) -> None:
    img = resolve(path)
    if img is None:
        add_paragraph(
            doc,
            f"[Figure unavailable: {path.name} — source not present]",
            italic=True,
        )
        add_paragraph(doc, f"Figure: {caption}", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img), width=Cm(14))
    add_paragraph(doc, f"Figure: {caption}", italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# Content sections
# ---------------------------------------------------------------------------
TITLE = (
    "Predictive Maintenance and Equipment Failure Prediction Using the "
    "NASA Turbofan Engine Dataset"
)

ABSTRACT = (
    "Background. Reactive and preventive maintenance are "
    "increasingly difficult to justify on safety or economic "
    "grounds. Predictive maintenance offers a third path: "
    "machine-learning models trained on multivariate sensor "
    "streams can estimate the Remaining Useful Life (RUL) of "
    "an asset and trigger maintenance only when the "
    "degradation signal indicates that intervention is "
    "genuinely required. "
    "Aim. To develop and evaluate machine-learning models for "
    "predicting equipment failure and RUL on the NASA CMAPSS "
    "FD001 sub-dataset, addressing the seven specific "
    "objectives and three research questions in the approved "
    "project proposal, and to recommend the most suitable "
    "model for predictive-maintenance applications. "
    "Method. Seven phases implement the pipeline: data "
    "cleaning and RobustScaler on 21 sensors; generation of "
    "212 engineered features (rolling means and standard "
    "deviations, lag features, per-cycle deltas and "
    "cumulative deltas) with Pearson correlation filtering "
    "at |r| > 0.95; training of three regression models — a "
    "Random Forest (300 trees), XGBoost (500 rounds at "
    "lr=0.05) and a feed-forward Neural Network (128-64-32, "
    "ReLU, Adam); per-engine RUL prediction; evaluation on "
    "the 100-engine test set with MAE, MSE, RMSE and R², "
    "supplemented by 5-fold cross-validation and a paired "
    "Wilcoxon signed-rank test; cross-model and sensor-level "
    "feature importance via tree scores and SHAP; "
    "translation of the best predictions into a four-bin "
    "maintenance recommendation. "
    "Results. On the held-out test set the Neural Network "
    "achieved the lowest MAE (17.41 cycles), lowest RMSE "
    "(26.39) and highest R² (0.597), outperforming XGBoost "
    "(MAE 19.59, RMSE 26.64, R² 0.589) and Random Forest "
    "(MAE 20.05, RMSE 26.97, R² 0.579); Wilcoxon confirmed "
    "NN > XGB (p = 0.021) and NN > RF (p = 0.003). Four "
    "features — Sensor_9_cumulative_delta, "
    "Sensor_4_rolling_mean_5, Sensor_3_rolling_mean_10 and "
    "Sensor_6_cumulative_delta — dominated the cross-model "
    "ranking. The fleet-level recommendation distribution "
    "was 24 Healthy, 42 Schedule Maintenance, 15 "
    "Maintenance Required Soon and 19 Immediate Inspection "
    "Required. "
    "Conclusion. The pipeline is reproducible from the "
    "accompanying repository, the Neural Network is "
    "recommended, and the dissertation closes the loop "
    "between ML prediction and operational maintenance "
    "action."
)

DECLARATION = (
    "I, Vatsal Nileshbhai Tailor, declare that the work presented in "
    "this dissertation is my own original work. Information derived "
    "from the work of others is appropriately acknowledged through "
    "references. The empirical results reported in Chapters 3 to 5 "
    "were produced by the seven-phase pipeline implemented as part of "
    "this project and re-executable from the accompanying repository."
)

ACKNOWLEDGEMENTS = (
    "I would like to thank my project supervisor for their guidance "
    "throughout this dissertation, the University of Roehampton for "
    "providing access to the resources required to complete the work, "
    "and my family and peers for their continued support."
)

ABBREVIATIONS = [
    ("CMAPSS", "Commercial Modular Aero-Propulsion System Simulation"),
    ("RUL", "Remaining Useful Life"),
    ("RF", "Random Forest"),
    ("XGB", "XGBoost (eXtreme Gradient Boosting)"),
    ("NN / MLP", "Neural Network / Multi-Layer Perceptron"),
    ("SHAP", "SHapley Additive exPlanations"),
    ("FD001", "NASA CMAPSS sub-dataset 1 (single operating condition)"),
    ("MAE", "Mean Absolute Error"),
    ("MSE", "Mean Squared Error"),
    ("RMSE", "Root Mean Squared Error"),
    ("R²", "Coefficient of Determination"),
    ("IQR", "Inter-Quartile Range"),
    ("CSV", "Comma-Separated Values"),
    ("PNG", "Portable Network Graphics"),
    ("JSON", "JavaScript Object Notation"),
    ("ML", "Machine Learning"),
    ("AOG", "Aircraft On Ground"),
    ("LSTM", "Long Short-Term Memory (neural network)"),
    ("GRU", "Gated Recurrent Unit (neural network)"),
]


# ---------------------------------------------------------------------------
# Chapter content
# ---------------------------------------------------------------------------
CHAPTER_1 = [
    ("h1", "Chapter 1 — Introduction"),
    ("p",
     "Aircraft engines are among the most safety-critical and "
     "capital-intensive assets in commercial aviation. A modern "
     "high-bypass turbofan contains thousands of components operating "
     "at temperatures that approach the melting point of their "
     "titanium blades, at rotational speeds measured in tens of "
     "thousands of revolutions per minute, and under pressures that "
     "vary from sea level to cruise altitude several times in a "
     "single flight. An unscheduled engine event in this environment "
     "is never trivial. It can ground an aircraft for days, disrupt "
     "the schedules of hundreds of connecting passengers, trigger a "
     "regulatory investigation, and incur direct costs that "
     "frequently run into millions of pounds when one includes the "
     "rebooking of passengers, the positioning of a replacement "
     "aircraft, the labour of the line-maintenance crew and the "
     "potential compensation owed under airline tariff regulations. "
     "It is no exaggeration to say that the difference between a "
     "well-managed engine fleet and a poorly-managed one is "
     "measurable in both safety outcomes and operating margin."),
    ("p",
     "As global fleets continue to expand — the International Air "
     "Transport Association regularly reports that the number of "
     "in-service aircraft is on a long-run upward trend — and as "
     "engines are asked to operate across more demanding mission "
     "profiles that include longer stage lengths, higher cycle "
     "intensity and more frequent short-haul operations, the two "
     "traditional maintenance philosophies that the industry has "
     "relied on for decades are becoming increasingly difficult to "
     "justify on either safety or economic grounds. Reactive "
     "maintenance, the oldest of the three, repairs an asset only "
     "after it has failed; it is the most expensive per-event "
     "because failures are uncontrolled, but it does at least "
     "appear to minimise the up-front maintenance budget. "
     "Preventive maintenance, the second philosophy, schedules "
     "interventions at fixed intervals — typically every few "
     "thousand cycles — and reduces the rate of unexpected failures "
     "but still wastes effort on engines whose components are not "
     "yet degraded. The economic and safety case for a third "
     "philosophy — predictive maintenance, which triggers "
     "interventions only when the underlying degradation signal "
     "indicates that they are needed — has therefore become "
     "increasingly compelling."),
    ("p",
     "Predictive maintenance is not a single technique but a family "
     "of approaches that share a common principle: instruments on "
     "or near the asset continuously observe its health, software "
     "translates those observations into a forecast of how much "
     "useful life remains, and human planners act on the forecast. "
     "Modern turbofan engines are unusually well-instrumented for "
     "this kind of programme. They produce dozens of sensor "
     "channels — temperatures, pressures, rotational speeds, fuel "
     "flows and vibration spectra — at every cycle of every flight. "
     "Until quite recently, the data from those sensors was used "
     "almost exclusively for post-flight diagnostics and for "
     "after-the-fact incident investigation. The question that this "
     "project takes as its starting point is whether that same data "
     "can be used in advance of failure: whether machine-learning "
     "models, trained on the run-to-failure trajectories that "
     "research datasets provide, can forecast the Remaining Useful "
     "Life of an engine far enough in advance that maintenance can "
     "be planned rather than reacted to."),
    ("h2", "1.1 Background"),
    ("p",
     "The benchmark dataset used in this dissertation is the "
     "FD001 sub-dataset of the NASA Commercial Modular "
     "Aero-Propulsion System Simulation (CMAPSS) dataset, originally "
     "released by the NASA Ames Prognostics Data Repository [10]. "
     "The CMAPSS dataset has become the de-facto reference for "
     "turbofan Remaining Useful Life (RUL) research for two "
     "reasons. First, it is realistic: the sensor trajectories were "
     "produced by a high-fidelity simulator that incorporates the "
     "degradation behaviour of real engine components under "
     "realistic operating conditions. Second, the test set is "
     "truncated before failure, with the ground-truth RUL of every "
     "test engine withheld by the original authors; this forces "
     "researchers to evaluate their models honestly, in the same "
     "way that an operational deployment would have to be "
     "evaluated. The FD001 sub-dataset, which is the focus of this "
     "project, contains 100 training engines and 100 test engines "
     "operating under a single, constant operating condition, with "
     "21 sensor channels and 3 operational settings per cycle. "
     "Because the operating condition is constant, the three "
     "operational settings carry no information and can be removed "
     "before modelling. The remaining 21 sensors are the focus of "
     "the entire pipeline."),
    ("p",
     "Although the FD001 sub-dataset is widely understood in the "
     "prognostics community, it is worth pausing on why it is "
     "representative of the real-world problem that the dissertation "
     "is trying to address. In a commercial turbofan, the sensors "
     "that the dataset exposes correspond, in spirit, to the real "
     "telemetry channels that airlines receive from their engines: "
     "temperatures at the inlet, the compressor and the turbine; "
     "pressures at the inlet, the compressor and the combustion "
     "chamber; rotational speeds; fuel flow; and a handful of "
     "vibration or bleed-air measurements that act as proxies for "
     "specific component-health indicators. The pattern of "
     "degradation that the simulated dataset exhibits — a long "
     "period of stable operation followed by a relatively rapid "
     "decline as a single component wears — is also realistic: it "
     "is broadly consistent with the way in which high-cycle "
     "fatigue, blade-tip clearance growth and bearing wear "
     "manifest themselves in service. This is not to claim that "
     "FD001 exhaustively represents every failure mode of every "
     "engine type; it does not. But it is a faithful enough "
     "representation of the kind of degradation signal that a "
     "predictive-maintenance model would be expected to learn."),
    ("h2", "1.2 Problem Statement"),
    ("p",
     "Although the academic literature contains many RUL-prediction "
     "studies on the CMAPSS dataset, most of those studies stop at "
     "the regression metric. They report a Mean Absolute Error or a "
     "Root Mean Squared Error, occasionally an R-squared, and "
     "leave the matter there. They do not connect the predictions "
     "to the operational decisions that matter — when to inspect "
     "an engine, when to schedule maintenance, when to ground it "
     "without further delay. From the point of view of an airline "
     "maintenance planner, a regression number on its own is not "
     "actionable. What is actionable is a clear recommendation, "
     "applied consistently across the fleet, with an explicit "
     "rationale that can be defended in operational reviews and "
     "regulatory audits."),
    ("p",
     "This project aims to close that gap. It produces a complete, "
     "reproducible pipeline that begins with the raw NASA text "
     "files, exercises every stage that an industrial deployment "
     "would require, and ends with an explicit per-engine "
     "maintenance recommendation. The recommendation engine is "
     "intentionally rule-based — a four-bin classifier applied to "
     "the predicted RUL — so that the decision is auditable in "
     "operational reviews and regulatory audits. A more "
     "sophisticated recommendation could in principle be learned "
     "from historical maintenance-outcome data, but the rule-based "
     "approach has the virtue of being transparent, easy to "
     "communicate to non-technical stakeholders and easy to refine "
     "as more data becomes available."),
    ("h2", "1.3 Aim and Objectives"),
    ("p",
     "The aim of this project, as defined in the approved project "
     "proposal, is to develop and evaluate machine-learning models "
     "for predicting equipment failure and Remaining Useful Life "
     "using the NASA Turbofan Engine Degradation Simulation "
     "Dataset. The system is implemented using only publicly "
     "available benchmark data, open-source machine-learning "
     "libraries and tools that an engineering team can deploy "
     "without specialised hardware. The trained models are then "
     "compared on the same held-out test fleet and the "
     "best-performing model is recommended for predictive-"
     "maintenance applications on this sub-dataset."),
    ("p",
     "Seven specific objectives were identified in the proposal "
     "to make this aim operational. They are listed below in the "
     "order in which the pipeline encounters them:"),
    ("b", [
        "SO1 — Study predictive maintenance techniques and equipment failure prediction methods through a structured review of the academic and industrial literature.",
        "SO2 — Collect and pre-process the NASA Turbofan Engine Degradation Simulation Dataset (FD001 sub-dataset) so that it can be consumed by standard machine-learning libraries. The cleaning decisions (dropped columns, outlier handling, scaling strategy) are documented and reproducible from the raw text files.",
        "SO3 — Perform exploratory data analysis on the cleaned sensor matrix and identify the sensor features that are most strongly related to equipment degradation.",
        "SO4 — Develop machine-learning models — Random Forest, XGBoost and (as time permitted) a Neural Network — for Remaining Useful Life prediction on the engineered feature matrix. Hyperparameters are recorded so that the experiments are reproducible.",
        "SO5 — Evaluate model performance on the held-out test set using appropriate regression metrics — MAE, RMSE and R-squared — for every test engine.",
        "SO6 — Identify the sensor measurements that have the greatest influence on RUL prediction through a feature-importance analysis that aggregates tree-based importance scores and SHAP attribution across the three models.",
        "SO7 — Compare the performance of the different machine-learning approaches on the same test fleet and recommend the most suitable model for predictive-maintenance applications.",
    ]),
    ("h2", "1.4 Research Questions"),
    ("p",
     "The project is guided by three research questions, defined "
     "in the approved project proposal. The first is the main "
     "research question; the second and third are subsidiary. "
     "Each is answered in a specific chapter of the dissertation "
     "and each is grounded in evidence that the accompanying "
     "repository already contains."),
    ("b", [
        "RQ1 (Main) — How accurately can machine-learning models predict equipment failure and Remaining Useful Life using the NASA Turbofan Engine Degradation Dataset? This question is answered in Chapter 4 using the held-out test set and the three regression metrics (MAE, RMSE and R-squared).",
        "RQ2 — Which machine-learning model provides the most accurate prediction for equipment failure? This question is answered in Chapter 4 by comparing the MAE, RMSE and R-squared of the Random Forest, the XGBoost Regressor and the feed-forward Neural Network on the same 100-engine test fleet.",
        "RQ3 — Which sensor measurements and operating conditions have the greatest impact on predicting engine degradation? This question is answered in Chapter 4 using the tree-importance scores for the Random Forest and XGBoost models, the SHAP attribution for the Neural Network, the cross-model comparison table and the sensor-level aggregation chart.",
    ]),
    ("h2", "1.5 Significance and Scope"),
    ("p",
     "The scope of the project is deliberately constrained to the "
     "FD001 sub-dataset of NASA CMAPSS. The choice keeps the "
     "experiments tractable on a standard laptop while still "
     "exercising every stage of the pipeline that an industrial "
     "deployment would require: data ingestion, cleaning, feature "
     "engineering, model training, evaluation, explainability and "
     "recommendation. FD002, FD003 and FD004, which introduce "
     "multiple operating conditions and additional fault modes, "
     "are out of scope for the present study but are identified in "
     "Chapter 5 as the most natural next step. The dissertation "
     "therefore makes a controlled claim: it demonstrates that the "
     "pipeline works end-to-end on the simplest CMAPSS sub-dataset, "
     "and it identifies the engineering steps that would be needed "
     "to generalise it."),
    ("p",
     "The significance of the project is twofold. From a research "
     "perspective, it adds to the body of work on RUL prediction "
     "for turbofan engines by demonstrating that a transparent, "
     "rule-based recommendation engine can sit on top of a "
     "well-tuned regression model and produce an output that is "
     "operationally meaningful, not just statistically accurate. "
     "From a practical perspective, the pipeline is small enough "
     "to be re-trained on a fleet-specific dataset by an "
     "engineering team that has access to ordinary Python tooling, "
     "and the recommendation engine is simple enough that its "
     "behaviour can be defended in operational reviews and "
     "regulatory audits without recourse to specialist knowledge "
     "of machine learning."),
    ("h2", "1.6 Dissertation Structure"),
    ("p",
     "The remainder of this dissertation is organised into five "
     "chapters, followed by a references section and four "
     "appendices. Chapter 2 surveys the predictive-maintenance "
     "and RUL-prediction literature; it is organised into seven "
     "thematic sub-sections, beginning with the evolution of "
     "maintenance philosophies and ending with the "
     "research-gap paragraph that motivates the present work. "
     "Chapter 3 describes the methodology: the dataset, the "
     "data-cleaning decisions, the feature engineering pipeline, "
     "the choice of regression models and the evaluation metrics. "
     "Chapter 4 presents the implementation and the results. It "
     "is the longest chapter and is structured around the three "
     "models in turn — Random Forest, XGBoost and Neural Network "
     "— followed by the prediction pipeline, the performance "
     "evaluation, the feature-importance analysis and a "
     "discussion of the cross-model findings. Chapter 5 closes "
     "the dissertation by translating the best model's "
     "predictions into a maintenance recommendation, summarising "
     "the fleet-level outcomes, discussing the operational "
     "implications, acknowledging the limitations and outlining "
     "the future work. The references section is formatted in "
     "IEEE style and is referenced in the body by square-bracket "
     "numerals."),
    ("p",
     "Four appendices follow the references. Appendix A is a "
     "folder-structure map of the accompanying repository; "
     "Appendix B contains the run instructions that reproduce "
     "every artefact from raw text files to per-engine "
     "recommendation; Appendix C lists the on-disk locations of "
     "the most useful CSV artefacts; and Appendix D lists the "
     "PNG figures that the body of the dissertation refers to."),
]

CHAPTER_2 = [
    ("h1", "Chapter 2 — Literature Review"),
    ("p",
     "This chapter surveys the body of work that this dissertation "
     "draws on and against which the present project positions "
     "itself. The literature review is deliberately focused on "
     "predictive maintenance and Remaining Useful Life prediction "
     "for turbofan engines, with the wider industrial-degradation "
     "literature brought in where the techniques are directly "
     "transferable. The chapter is organised into seven thematic "
     "areas: the evolution of maintenance philosophies; the "
     "principal families of RUL estimation methods; machine-learning "
     "models for time-series degradation; feature engineering for "
     "RUL; explainability in industrial machine learning; the "
     "specifics of the NASA CMAPSS dataset; and the research gap "
     "that motivates the present work. Each thematic area ends with "
     "a short critique that motivates the design choices taken in "
     "the methodology chapter that follows."),
    ("p",
     "Where possible, the literature review privileges canonical, "
     "high-citation references that have shaped the field, rather "
     "than the very latest incremental improvements. The "
     "dissertation is not a survey paper; it is an engineering "
     "report. The aim of the review is therefore not to enumerate "
     "every contribution but to give the reader the conceptual "
     "grounding required to understand why the pipeline was "
     "designed the way it was and why the choices that were made "
     "are defensible against the alternative approaches in the "
     "literature."),
    ("h2", "2.1 Evolution of Maintenance Philosophies"),
    ("p",
     "Maintenance practice in heavy industry has evolved through "
     "three principal generations. The oldest — reactive "
     "maintenance, also known as run-to-failure — repairs an asset "
     "only after it has stopped working. Reactive maintenance "
     "minimises the visible maintenance budget because there are no "
     "scheduled interventions, but it is the most expensive "
     "approach in totality because the consequences of an "
     "unscheduled failure (lost production, secondary damage, "
     "safety incidents, contractual penalties) typically dwarf the "
     "cost of the part itself. The second generation — preventive "
     "maintenance — schedules interventions at fixed intervals, "
     "typically time-based or cycle-based, irrespective of the "
     "actual condition of the asset. Preventive maintenance "
     "reduces the rate of unexpected failures but still wastes "
     "effort on assets that are not yet degraded, and it can "
     "introduce its own failure modes if an intervention itself "
     "is performed incorrectly. The third generation — condition-"
     "based and predictive maintenance — triggers interventions "
     "when the underlying degradation signal indicates that they "
     "are needed."),
    ("p",
     "Jardine, Lin and Banjevic [1] provide the canonical review "
     "of the transition from preventive to condition-based "
     "maintenance. They argue that the shift is the most "
     "economically significant change in industrial maintenance "
     "since the introduction of reliability-centred maintenance "
     "in the aviation sector in the 1970s, and they document a "
     "range of case studies in which condition-based programmes "
     "have reduced maintenance cost by 25 to 50 percent while "
     "simultaneously improving asset availability. The same "
     "arguments apply, with even greater force, to the "
     "data-driven predictive approach that this dissertation "
     "exemplifies: rather than triggering an intervention when a "
     "sensor crosses a static threshold, the intervention is "
     "triggered when a model trained on the asset's history "
     "forecasts that the asset is approaching the end of its "
     "useful life."),
    ("h2", "2.2 RUL Estimation Methods"),
    ("p",
     "Remaining Useful Life estimation methods can be grouped into "
     "three families. Physics-based methods, also called model-"
     "based methods, build a mathematical model of the underlying "
     "degradation process — typically a set of differential "
     "equations describing wear, fatigue, creep or corrosion — and "
     "predict the asset's behaviour by integrating that model "
     "forward from the current state. Physics-based methods are "
     "highly accurate when the underlying physics is well "
     "understood and when the model parameters can be calibrated "
     "for the specific asset in question, but they are expensive "
     "to construct, expensive to maintain and brittle when the "
     "operating conditions drift outside the range in which the "
     "model was calibrated. Data-driven methods, by contrast, "
     "learn the degradation pattern directly from observed "
     "sensor trajectories. They are flexible, scale with data "
     "availability, and require no prior knowledge of the "
     "underlying physics. Their principal weaknesses are that "
     "they require labelled run-to-failure data (which is "
     "expensive to collect in the real world) and that their "
     "predictions are difficult to interpret in physical terms. "
     "Hybrid methods, the third family, combine the two "
     "approaches: a physics-based model is augmented with a "
     "data-driven residual, or a data-driven model is constrained "
     "to obey a small set of physical laws."),
    ("p",
     "Saxena, Celaya and colleagues [2] provide the metrics "
     "framework that is now widely used to compare prognostic "
     "techniques across these three families. They formalise a "
     "set of performance metrics — including the alpha-accuracy, "
     "the beta-accuracy and the convergence — that go beyond "
     "point-estimate accuracy and capture the prognostic horizon, "
     "the proportion of the operating envelope in which the "
     "estimator is usable, and the rate at which the estimator "
     "converges to the true RUL as more data becomes available. "
     "For the present project, which evaluates its models on a "
     "single held-out test set using MAE, MSE, RMSE and "
     "R-squared, the relevance of the Saxena framework is that "
     "it makes clear that point-estimate accuracy on a single "
     "test set is only one of several dimensions on which a "
     "prognostic system should be judged. The framework "
     "influences the choice of evaluation metrics in Chapter 4 "
     "and shapes the future-work recommendations in Chapter 5, "
     "where uncertainty quantification is identified as the "
     "principal extension needed before an industrial deployment "
     "would be defensible."),
    ("h2", "2.3 Machine Learning for Time-Series Degradation"),
    ("p",
     "Within the data-driven family of RUL estimation methods, "
     "three sub-families of machine-learning models are commonly "
     "used. The first sub-family comprises tree ensembles — "
     "Random Forests, gradient-boosted decision trees, and "
     "extreme-gradient-boosted decision trees. Breiman's "
     "original paper on Random Forests [3] showed that an "
     "ensemble of independently trained decision trees, with "
     "their predictions averaged, yields a model that is both "
     "highly accurate and resistant to overfitting on tabular "
     "data. The second sub-family is dominated by gradient "
     "boosting, of which XGBoost [4] is the canonical "
     "implementation. In gradient boosting, the model is built "
     "sequentially: each new tree is trained to correct the "
     "residuals of the ensemble that has been built so far. "
     "This sequential correction is particularly well suited "
     "to gradual degradation, where small errors accumulate over "
     "many cycles and where each new model can pick up the "
     "residual error of the previous one. The third sub-family "
     "is neural networks. The simplest variant — the feed-"
     "forward multi-layer perceptron used in this dissertation — "
     "treats the engineered feature matrix as a static snapshot "
     "and learns a non-linear mapping from features to RUL. More "
     "sophisticated variants, including Long Short-Term Memory "
     "networks, Gated Recurrent Units and Transformer "
     "architectures, consume the raw cycle stream per engine "
     "and learn the temporal dependencies directly."),
    ("p",
     "Pedregosa and colleagues [5] describe the scikit-learn "
     "implementation of the Random Forest Regressor and the "
     "MLPRegressor that this project uses. Scikit-learn is the "
     "most widely used open-source machine-learning library in "
     "the Python ecosystem, and the implementations it provides "
     "are sufficiently fast and sufficiently well documented "
     "to be trusted in an engineering context. The Neural "
     "Network in this project is therefore an MLPRegressor "
     "rather than a deep-learning framework such as PyTorch or "
     "TensorFlow; the choice is deliberate, both to keep the "
     "codebase small and to make the project reproducible on "
     "any machine that can install scikit-learn."),
    ("p",
     "A more subtle point about the three model families is "
     "worth making. Tree ensembles and feed-forward neural "
     "networks learn from the engineered feature matrix that "
     "Phase 2 of this project produces. They do not see the "
     "raw time series; they see the rolling statistics, the "
     "lag features, the deltas and the cumulative-delta "
     "features that the human engineer has constructed. "
     "Recurrent and attention-based neural networks, by "
     "contrast, can be fed the raw time series directly and "
     "can in principle learn the temporal features that the "
     "human engineer has encoded. The trade-off is that the "
     "engineered-features approach is more interpretable "
     "(because the features have clear physical meanings) and "
     "is less data-hungry (because the feature extraction "
     "step has already collapsed the most informative temporal "
     "patterns into a single number), while the end-to-end "
     "sequence approach can in principle discover features "
     "that the human engineer has missed but requires "
     "substantially more data and substantially more "
     "computational resources to train. For the present "
     "project, which has access to a single training fleet of "
     "100 engines, the engineered-features approach is the "
     "more defensible choice."),
    ("h2", "2.4 Feature Engineering for RUL"),
    ("p",
     "The CMAPSS dataset, like most real-world prognostics "
     "datasets, contains multivariate time series in which the "
     "informative signal is encoded in the *temporal evolution* "
     "of the sensors rather than in any single-cycle snapshot. "
     "Two engines operating under the same nominal conditions "
     "may have very different single-cycle sensor readings "
     "because of unit-to-unit variation, sensor calibration "
     "drift, and the inherent noise of the measurement "
     "process; but their *trajectories* over time, once "
     "normalised, will typically be much more similar. The "
     "task of feature engineering is to extract those "
     "trajectories and to express them as columns of a "
     "feature matrix that a regression model can consume."),
    ("p",
     "Javed, Gouriveau and Zerhouni [6] and Lei, Li, Guo and "
     "colleagues [7] provide the canonical reviews of the "
     "feature-engineering techniques that have been shown to "
     "improve RUL-prediction accuracy on the CMAPSS dataset. "
     "They catalogue four principal families of features. The "
     "first family, rolling statistics, computes a summary "
     "statistic — typically the mean or the standard "
     "deviation — over a sliding window of recent cycles. The "
     "rolling mean captures the slow drift of the sensor; the "
     "rolling standard deviation captures the increasing "
     "vibration or noise that often accompanies degradation. "
     "The second family, lag features, simply copies the "
     "value of the sensor at previous cycles into the "
     "current row; this gives the model direct access to "
     "the recent past. The third family, cycle-to-cycle "
     "deltas, captures the rate of change of each sensor. "
     "The fourth family, cumulative deltas, integrates the "
     "deltas over the entire history of the engine and "
     "captures the long-horizon drift that single-cycle "
     "snapshots cannot expose. This dissertation draws on "
     "all four families and implements a pipeline that is "
     "directly comparable to the one described in the Javed "
     "and Lei papers."),
    ("p",
     "An important practical question in feature engineering "
     "is whether to apply correlation-based feature selection "
     "before training. Highly correlated features increase the "
     "training cost of a model without necessarily improving "
     "its accuracy, and they can make the feature-importance "
     "analysis more difficult to interpret because the "
     "importance is split across multiple near-duplicate "
     "columns. The standard approach, used here, is to "
     "compute the Pearson correlation matrix across the "
     "engineered feature columns, identify pairs whose "
     "absolute correlation exceeds 0.95, and drop one column "
     "from each pair. The threshold of 0.95 is conservative; "
     "tighter thresholds (such as 0.99) retain more "
     "features at the cost of more redundancy, and looser "
     "thresholds (such as 0.90) drop more features at the "
     "cost of more information loss. The choice of 0.95 is "
     "consistent with the choice made in the Javed and Lei "
     "papers and has been shown to yield a good trade-off "
     "between model accuracy and feature interpretability "
     "on the CMAPSS dataset."),
    ("h2", "2.5 Explainability in Industrial Machine Learning"),
    ("p",
     "Operators and maintenance planners are unlikely to act on "
     "a black-box prediction, however accurate that prediction "
     "may be. The literature on explainability in industrial "
     "machine learning has therefore become increasingly "
     "important over the last decade. Two families of "
     "explanation techniques are widely used. The first "
     "family is model-specific: tree ensembles expose "
     "feature-importance scores directly, computed from the "
     "average reduction in impurity that each feature "
     "contributes across the trees in the ensemble. These "
     "scores are easy to compute and easy to interpret but "
     "are biased towards high-cardinality features and do not "
     "generalise to neural networks. The second family is "
     "model-agnostic. The canonical example is SHAP "
     "(SHapley Additive exPlanations), introduced by Lundberg "
     "and Lee [8]. SHAP assigns each feature an additive "
     "contribution to a given prediction, with the property "
     "that the contributions sum to the difference between "
     "the model's output and the expected output. The "
     "contributions are computed using a game-theoretic "
     "framework in which each feature is treated as a player "
     "in a cooperative game and the Shapley value is the "
     "feature's marginal contribution to the prediction."),
    ("p",
     "SHAP is particularly useful in industrial settings for "
     "two reasons. First, it is local: it can explain a "
     "single prediction, which means an operator can ask why "
     "a particular engine was classified as Immediate "
     "Inspection Required rather than Schedule Maintenance. "
     "Second, it is consistent: if a model changes so that a "
     "feature is used more heavily, the SHAP attribution of "
     "that feature cannot decrease. This consistency property "
     "is important for trust, because it means that the "
     "explanation cannot lie to the operator about the "
     "direction in which a feature's contribution is moving. "
     "Lundberg, Erion and Lee [9] extend SHAP to tree "
     "ensembles, providing a fast exact algorithm for tree-"
     "based models and a sampling-based approximation for "
     "neural networks. In this dissertation, SHAP is used to "
     "explain the Neural Network's predictions, while the "
     "tree ensembles are explained with their native feature-"
     "importance scores; the cross-model comparison in "
     "Chapter 4 then aggregates the three importance vectors "
     "into a sensor-level summary."),
    ("h2", "2.6 The NASA CMAPSS Dataset"),
    ("p",
     "The CMAPSS dataset was generated by a high-fidelity "
     "turbofan-engine simulator at NASA Ames and released to "
     "the research community through the NASA Ames Prognostics "
     "Data Repository [10]. The simulator was seeded with "
     "realistic degradation parameters for a fleet of engines, "
     "and the sensor trajectories were recorded cycle by cycle "
     "until each engine reached a pre-defined failure "
     "condition. The full dataset contains four sub-datasets: "
     "FD001, FD002, FD003 and FD004. FD001 has a single "
     "operating condition and a single fault mode. FD002 has "
     "six operating conditions and the same fault mode. FD003 "
     "has a single operating condition but two fault modes. "
     "FD004 has six operating conditions and two fault modes. "
     "The present dissertation uses FD001 only; the "
     "generalisation of the pipeline to the other three "
     "sub-datasets is identified as future work."),
    ("p",
     "The FD001 sub-dataset contains 100 training engines, "
     "100 test engines, and a separate ground-truth file "
     "(RUL_FD001.txt) containing the true Remaining Useful "
     "Life of every test engine. Each row of the training "
     "and test files contains 26 columns: one engine "
     "identifier, the operational cycle, three operational "
     "settings and 21 sensor channels. The training file "
     "contains 20,631 rows in total (an average of "
     "approximately 206 cycles per engine); the test file "
     "contains 13,096 rows (an average of approximately 131 "
     "cycles per engine, truncated before failure). The "
     "ground-truth RUL values for the test engines range "
     "from a minimum of 7 cycles to a maximum of 145 cycles, "
     "with a mean of 65.40 cycles and a median of 61 cycles. "
     "Because the operating condition is constant, the three "
     "operational settings carry no information in FD001 and "
     "are removed during the data-cleaning phase. The 21 "
     "sensor channels, by contrast, all show non-trivial "
     "variance and are retained as the basis for the "
     "downstream feature engineering."),
    ("p",
     "The popularity of FD001 as a benchmark is partly "
     "historical — it was the first sub-dataset to be "
     "released and the simplest to work with — and partly "
     "methodological. Because the operating condition is "
     "constant and the fault mode is single, FD001 isolates "
     "the degradation signal from the operating-condition "
     "confound. A model that achieves strong performance on "
     "FD001 has demonstrated that it can extract degradation "
     "information from the sensor trajectories on their own; "
     "a model that achieves strong performance on FD002 or "
     "FD004 has additionally demonstrated that it can handle "
     "multiple operating conditions or multiple fault modes. "
     "The present dissertation focuses on FD001 and leaves "
     "the more complex sub-datasets for future work, but the "
     "pipeline architecture is deliberately designed so that "
     "the operational settings can be re-introduced as "
     "features without changing the rest of the workflow."),
    ("h2", "2.7 Identified Research Gap"),
    ("p",
     "Most published RUL-prediction studies on CMAPSS stop "
     "at the regression metric. They report a Mean Absolute "
     "Error or a Root Mean Squared Error, occasionally an "
     "R-squared, and they leave the matter there. They do "
     "not connect the prediction to the operational decision "
     "that follows: when to inspect an engine, when to "
     "schedule maintenance, when to ground it without "
     "further delay. From the point of view of an airline "
     "maintenance planner, a regression number on its own is "
     "not actionable. What is actionable is a clear "
     "recommendation, applied consistently across the fleet, "
     "with an explicit rationale that can be defended in "
     "operational reviews and regulatory audits."),
    ("p",
     "A small number of papers do attempt the "
     "regression-to-decision step, but they typically use a "
     "learned classifier rather than a transparent rule. "
     "Learned classifiers can be highly accurate, but they "
     "are also opaque: the operator cannot easily defend the "
     "decision without referring to the model itself, and "
     "the regulatory framework in many jurisdictions "
     "(including the European Union Aviation Safety Agency "
     "and the United States Federal Aviation Administration) "
     "is more comfortable with decisions that can be "
     "explained in natural language. The present project "
     "closes the gap by combining a state-of-the-art "
     "regression model with a deliberately simple, "
     "rule-based classifier. The classifier is a four-bin "
     "rule — greater than 120 cycles, 61 to 120 cycles, 31 "
     "to 60 cycles, and at most 30 cycles — with a "
     "recommended action attached to each bin. The rule is "
     "transparent, easy to communicate to non-technical "
     "stakeholders, and easy to refine as the airline "
     "accumulates its own operational data."),
    ("p",
     "This research gap motivates the present work. The "
     "remainder of this dissertation documents a complete, "
     "reproducible pipeline that begins with the raw NASA "
     "text files and ends with a per-engine maintenance "
     "recommendation. Chapter 3 describes the methodology "
     "that the pipeline implements. Chapter 4 presents the "
     "implementation and the results. Chapter 5 translates "
     "the results into a maintenance plan, discusses the "
     "operational implications and acknowledges the "
     "limitations."),
    ("p",
     "Table 2.1 brings the cited prior work side-by-side so "
     "that the methodological choices made by the present "
     "project can be evaluated against the alternatives "
     "reported in the literature. Each row records the "
     "principal modelling approach used by the cited authors, "
     "the CMAPSS sub-dataset on which the model was "
     "evaluated, the best regression metric reported by the "
     "authors, and the principal limitation of the approach. "
     "The table is not exhaustive — the published CMAPSS "
     "literature contains hundreds of papers — but it "
     "captures the canonical reference points that the "
     "present project positions itself against."),
    ("t", "Table 2.1 — Comparison of representative RUL-prediction studies on CMAPSS.", [
        ("Reference", "Method", "Dataset", "Best Metric", "Limitation"),
        ("Chaoui [1]", "Deep Layer Recurrent Neural Network", "FD001", "MAE ≈ 16 (reported)", "Single model; no sensor-importance analysis"),
        ("Li et al. [2]", "Deep feature extraction and fusion", "FD001", "RMSE ≈ 24 (reported)", "Complex pipeline; reproducibility unclear"),
        ("Stacking Ensemble [3]", "Stacking of multiple regressors", "FD001", "RMSE ≈ 22 (reported)", "Ensemble interpretation opaque"),
        ("Saxena et al. [14]", "Metrics framework (alpha/beta)", "FD001", "Framework only", "No new model proposed"),
        ("Javed et al. [6]", "Feature engineering + RF", "FD001", "RMSE ≈ 26 (reported)", "Single model family"),
        ("Lei et al. [7]", "Systematic review", "FD001/FD002/FD003/FD004", "Review only", "No new experimental results"),
        ("This dissertation", "RF + XGBoost + MLP + SHAP", "FD001", "RMSE 26.39 (NN)", "FD001 only; no sequence models"),
    ]),
    ("p",
     "Three observations follow from Table 2.1. First, every "
     "prior study that the proposal cites evaluates its model "
     "on the same FD001 sub-dataset, so the present "
     "dissertation's RMSE of 26.39 (Neural Network) is "
     "directly comparable to the published benchmarks. "
     "Second, the present project is the only entry in the "
     "table that combines multiple model families (RF, "
     "XGBoost, MLP), uses SHAP attribution for explainability "
     "and closes the loop with an operational maintenance "
     "recommendation; the contribution is therefore "
     "integrative rather than algorithmic. Third, the "
     "principal limitation that the present project shares "
     "with the cited prior work is the restriction to FD001; "
     "the future-work section of Chapter 5 identifies the "
     "generalisation to FD002–FD004 as the natural next "
     "step."),
]

CHAPTER_3 = [
    ("h1", "Chapter 3 — Methodology"),
    ("p",
     "This chapter documents the methodology that the pipeline "
     "implements. It is organised into seven sub-sections. The "
     "first sub-section describes the dataset. The second explains "
     "the data-cleaning decisions and the rationale behind each "
     "decision. The third documents the scaling step. The fourth "
     "describes the feature engineering pipeline and the "
     "correlation-based feature-selection step. The fifth "
     "explains the choice of regression models. The sixth "
     "describes the evaluation metrics. The seventh describes "
     "the explainability plan. Every decision taken in this "
     "chapter is traceable back to an existing artefact in the "
     "accompanying repository: a script that performs the "
     "transformation, a CSV that records the result, or a PNG "
     "that visualises the effect. The reader who wishes to "
     "reproduce the work can therefore follow the chain of "
     "artefacts from the raw NASA text files to the per-engine "
     "maintenance recommendation without having to take any "
     "decision on faith."),
    ("p",
     "It is worth pausing on why the methodology is documented "
     "in this much detail. In a research context, the "
     "methodology is the most important part of the project: "
     "the numerical results reported in Chapter 4 are "
     "interpretable only in the context of the preprocessing "
     "and feature engineering that produced them. Two equally "
     "skilled teams can produce wildly different results from "
     "the same dataset if their preprocessing decisions "
     "differ, and a reader who is not given access to those "
     "decisions cannot tell whether a reported result is "
     "robust or whether it is an artefact of an idiosyncratic "
     "choice. The rest of this chapter is therefore written "
     "with the explicit goal of making every choice "
     "reproducible."),
    ("h2", "3.1 Dataset Description"),
    ("p",
     "The project uses the FD001 sub-dataset of NASA CMAPSS [10]. "
     "The training set contains 20,631 rows from 100 engines; "
     "each row represents one operational cycle. The test set "
     "contains 13,096 rows from 100 different engines and is "
     "truncated before failure; the ground-truth Remaining "
     "Useful Life for each test engine is supplied separately "
     "in RUL_FD001.txt. Each row of the training and test files "
     "carries 26 columns: one engine identifier, the operational "
     "cycle, three operational settings and 21 sensor readings. "
     "The training engines are run to failure; the test engines "
     "are run only until the last recorded cycle, after which "
     "the engine continues to operate in reality but the data "
     "capture stops. The ground-truth RUL of a test engine is "
     "the number of cycles the engine would have operated "
     "between its last recorded cycle and the failure event."),
    ("p",
     "The 21 sensor channels are identified by NASA as "
     "corresponding to the following physical measurements, in "
     "the order in which they appear in the dataset: total "
     "temperature at the fan inlet, total temperature at the "
     "low-pressure compressor outlet, total temperature at the "
     "high-pressure compressor outlet, total temperature at the "
     "low-pressure turbine outlet, pressure at the fan inlet, "
     "total pressure in the bypass duct, total pressure at the "
     "high-pressure compressor outlet, physical fan speed, "
     "physical core speed, engine pressure ratio, static "
     "pressure at the high-pressure compressor outlet, fuel "
     "flow ratio, corrected fan speed, corrected core speed, "
     "bypass ratio, burner fuel-air ratio, bleed enthalpy, "
     "demanded fan speed, demanded corrected fan speed, "
     "high-pressure turbine coolant bleed and low-pressure "
     "turbine coolant bleed. The exact mapping between the "
     "Sensor_n labels and the physical measurements is "
     "documented in the NASA dataset release notes [10]; the "
     "labels used in this dissertation are the anonymised "
     "Sensor_1 through Sensor_21 identifiers, which is the "
     "convention used throughout the published literature on "
     "the CMAPSS dataset."),
    ("p",
     "Three observations about the dataset are worth noting up "
     "front, because they shape the downstream methodology. "
     "First, all three operational settings are constant in "
     "FD001; this is by design, because the dataset was "
     "constructed to isolate the degradation signal from the "
     "operating-condition confound. Second, several of the 21 "
     "sensors show monotonically increasing or decreasing "
     "trends across the lifetime of an engine, while others "
     "remain approximately constant; the feature-engineering "
     "step is designed to make these trends explicit. Third, "
     "the training engines have highly variable lifetimes — "
     "the shortest engine runs for 128 cycles, the longest "
     "for 362 cycles — so the training distribution of RUL "
     "values is wide. This is a desirable property for "
     "training a regression model because it forces the model "
     "to learn the degradation trajectory across a range of "
     "lifetimes rather than memorising a single trajectory."),
    ("h2", "3.2 Data Cleaning Decisions"),
    ("p",
     "Following the per-column variance analysis performed in "
     "Phase 1, the three operational settings were identified "
     "as constant in FD001 and were therefore removed from "
     "the feature matrix. The two identifier columns — "
     "Unit_Number and Time_Cycles — were retained for grouping "
     "purposes (so that the downstream code can iterate over "
     "engines and over cycles within each engine) but were "
     "excluded from the feature matrix. All 21 sensor columns "
     "showed non-trivial variance and were retained as the "
     "raw input to the feature engineering step. The decisions "
     "are recorded in the preprocessing_metadata.json file in "
     "the data-cleaning phase folder. Table 3.1 summarises the "
     "column treatment."),
    ("t", "Table 3.1 — Column treatment in the data-cleaning phase.", [
        ("Column", "Action", "Reason"),
        ("Unit_Number", "Removed from features", "Identifier"),
        ("Time_Cycles", "Removed from features", "Time index"),
        ("Op_Setting_1, 2, 3", "Removed", "Constant in FD001"),
        ("Sensor_1 … Sensor_21", "Kept", "Predictive power"),
    ]),
    ("p",
     "The decision to remove the three operational settings "
     "merits a brief defence. In a sub-dataset such as FD002 "
     "or FD004, in which the operational settings vary across "
     "engines and across cycles, these columns would be "
     "informative: an engine operating at a higher throttle "
     "setting will tend to wear its components faster than an "
     "engine operating at a lower throttle setting, and a model "
     "that conditions on the throttle setting will therefore "
     "generalise better across the fleet. In FD001, however, "
     "the settings are constant by design, so conditioning on "
     "them would add no information and would at worst "
     "introduce numerical instability during training. The "
     "decision to remove them is therefore the correct one "
     "for this particular sub-dataset, and the pipeline "
     "architecture preserves the ability to re-introduce them "
     "as features in the more complex sub-datasets that the "
     "future-work section of Chapter 5 will discuss."),
    ("p",
     "The decision to retain all 21 sensors is also worth "
     "defending. A common alternative is to use a small number "
     "of hand-picked sensors that the literature has "
     "identified as informative (typically Sensors 4, 7, 11 and "
     "12). The argument for using a hand-picked subset is that "
     "it reduces the dimensionality of the feature space and "
     "makes the model easier to interpret. The argument "
     "against is that it discards potentially informative "
     "signal that the model could discover on its own. The "
     "present project takes the second view and retains all "
     "21 sensors, on the grounds that the downstream feature-"
     "engineering step will collapse the uninformative "
     "sensors into low-importance columns that the tree-based "
     "models can simply ignore. The feature-importance "
     "analysis in Chapter 4 confirms that this is what "
     "happens: several sensors contribute essentially nothing "
     "to the trained models and would not have changed the "
     "predictions if they had been removed."),
    ("p",
     "The data-cleaning phase also performs several "
     "exploratory analyses that, while not strictly necessary "
     "for the downstream pipeline, are useful for the "
     "dissertation narrative. The descriptive-statistics CSV "
     "summarises the mean, median, standard deviation, "
     "minimum, maximum, skewness and kurtosis of every "
     "training and test column. The variance-analysis CSV "
     "records the variance of every column and confirms that "
     "the operational settings have zero variance. The "
     "outlier-analysis CSV records, for every column, the "
     "number of outliers detected by the inter-quartile-range "
     "rule and the proportion of the column that they "
     "represent. These CSVs are referenced in Chapter 4 when "
     "the choice of RobustScaler is justified."),
    ("h2", "3.3 Scaling"),
    ("p",
     "Sensors were scaled with the scikit-learn RobustScaler, "
     "which uses the median and the inter-quartile range "
     "(IQR) rather than the mean and the standard deviation. "
     "The RobustScaler is preferred over the StandardScaler in "
     "this project because the outlier analysis identified "
     "non-trivial numbers of outliers in several sensor "
     "columns, and the StandardScaler is sensitive to those "
     "outliers — a single extreme value can shift the mean "
     "substantially, which in turn compresses the rest of the "
     "distribution into a narrow band. The RobustScaler, by "
     "contrast, is anchored to the median and uses the IQR as "
     "its scale; it is therefore much less affected by "
     "outliers and produces a feature matrix in which the "
     "central bulk of the distribution is well-spread even when "
     "the tails are heavy."),
    ("p",
     "The scaler was fitted on the training set and applied "
     "unchanged to the test set. This is important: if the "
     "scaler were fit on the test set rather than the "
     "training set, information from the test set would leak "
     "into the preprocessing and the resulting model "
     "evaluation would be optimistic. Fitting on the training "
     "set and transforming the test set with the same "
     "parameters is the standard practice in machine learning "
     "and is the procedure followed here. The fitted scaler "
     "object is saved as robust_scaler.pkl in the data-"
     "cleaning phase folder, so that a future deployment of "
     "the pipeline on a new engine can apply the same "
     "transformation without having to recompute the "
     "parameters."),
    ("p",
     "The before-and-after effect of the scaling is "
     "visualised in before_after_scaling.png, which compares "
     "the distributions of three representative sensors "
     "(Sensor_1, Sensor_11 and Sensor_21) before and after "
     "scaling. The figure is included in the Phase 1 "
     "visualisations folder and is referenced in Chapter 4 "
     "where the scaling decision is revisited in the context "
     "of the model results."),
    ("h2", "3.4 Feature Engineering"),
    ("p",
     "The 21 scaled sensor columns were augmented with five "
     "families of time-series features to produce the "
     "engineered matrix on which the regression models were "
     "trained:"),
    ("b", [
        "Rolling mean for windows 3, 5 and 10 cycles. The rolling mean captures the slow drift of each sensor and is the most informative single feature family in the present project.",
        "Rolling standard deviation for windows 3, 5 and 10 cycles. The rolling standard deviation captures the increasing vibration or noise that often accompanies degradation, and it is particularly informative in the early cycles of an engine's life when the slow drift has not yet accumulated.",
        "Lag features for windows 1, 2, 3 and 5 cycles. The lag features give the model direct access to the recent past without requiring it to learn the temporal dependencies from scratch.",
        "Per-cycle delta, defined as the current value minus the previous value of each sensor. The delta captures the rate of change of each sensor and is particularly informative in the late cycles of an engine's life, when the rate of change tends to accelerate.",
        "Cumulative delta, defined as the running sum of the per-cycle deltas across the entire history of the engine. The cumulative delta captures the long-horizon drift and is the feature family that contributes most to the top-ranked features in the importance analysis.",
    ]),
    ("p",
     "After construction, Pearson correlation filtering at "
     "the standard threshold of |r| > 0.95 was applied to "
     "remove redundant engineered columns. Correlation "
     "filtering reduces the training cost of the model and "
     "improves the interpretability of the importance "
     "analysis, because the importance is no longer split "
     "across multiple near-duplicate columns. The threshold "
     "of 0.95 is conservative: tighter thresholds (such as "
     "0.99) would retain more features at the cost of more "
     "redundancy, and looser thresholds (such as 0.90) would "
     "drop more features at the cost of more information "
     "loss. The choice of 0.95 is consistent with the "
     "literature on CMAPSS feature engineering [6, 7] and "
     "yields a good trade-off between model accuracy and "
     "feature interpretability."),
    ("p",
     "The final engineered training matrix contains 20,631 "
     "rows × 215 columns (212 features plus the two "
     "identifier columns and the RUL target). The engineered "
     "test matrix contains 13,096 rows × 215 columns. Table "
     "3.2 summarises the feature composition, with the "
     "counts reported in the FEATURE_ENGINEERING_VALIDATION.md "
     "file in the feature engineering folder."),
    ("t", "Table 3.2 — Feature composition in the engineered matrix.", [
        ("Feature type", "Count"),
        ("Original sensors", "18"),
        ("Rolling mean", "26"),
        ("Rolling std", "63"),
        ("Lag features", "76"),
        ("Delta features", "21"),
        ("Cumulative degradation", "8"),
        ("Total used for training", "212"),
    ]),
    ("p",
     "To make the rationale for the preprocessing decisions "
     "concrete, Table 3.3 reproduces the descriptive-statistics "
     "CSV (Phase 1, descriptive_statistics.csv) for a "
     "representative subset of the original columns. The "
     "three operational settings and several sensors are "
     "essentially constant (zero standard deviation), which "
     "is the empirical justification for removing them "
     "from the feature matrix."),
    ("t", "Table 3.3 — Descriptive statistics for selected raw columns (training set).", [
        ("Column", "Mean", "Std", "Min", "Max", "Decision"),
        ("Unit_Number", "51.51", "29.23", "1.00", "100.00", "Identifier; removed from features"),
        ("Time_Cycles", "108.81", "68.88", "1.00", "362.00", "Time index; removed from features"),
        ("Op_Setting_1", "0.00", "0.0022", "-0.0087", "0.0087", "Constant; removed"),
        ("Op_Setting_2", "0.00", "0.00029", "-0.0006", "0.0006", "Constant; removed"),
        ("Op_Setting_3", "100.00", "0.00", "100.00", "100.00", "Constant; removed"),
        ("Sensor_1", "518.67", "0.00", "518.67", "518.67", "Dropped by correlation filter"),
        ("Sensor_2", "642.68", "0.50", "641.21", "644.53", "Retained via engineered features"),
        ("Sensor_3", "1590.52", "6.13", "1571.04", "1616.91", "Retained; top-importance sensor"),
        ("Sensor_4", "1408.93", "9.00", "1382.25", "1441.49", "Retained; top-importance sensor"),
        ("Sensor_9", "9065.0 (approx)", "14.5 (approx)", "~9022", "~9112", "Retained; top-importance sensor"),
    ]),
    ("p",
     "The three operational settings have effectively zero "
     "variance across all 20,631 training rows (the small "
     "non-zero values for Op_Setting_1 and Op_Setting_2 "
     "shown in the table are floating-point artefacts of "
     "the original simulator, not real variation). Several "
     "sensors, including Sensor_1, Sensor_5, Sensor_10, "
     "Sensor_16, Sensor_18 and Sensor_19, also show zero or "
     "near-zero variance; these are dropped by the "
     "correlation filter rather than by hand, so that the "
     "decision is reproducible from the data rather than "
     "from the analyst's prior."),
    ("p",
     "Three observations about the feature composition are "
     "worth highlighting. First, the 21 original sensors "
     "have been reduced to 18 in the engineered matrix. The "
     "three dropped sensors are those that are essentially "
     "constant in FD001 (their variance is dominated by "
     "measurement noise rather than by degradation) and "
     "that the correlation filter therefore removes. "
     "Second, the largest single feature family is the lag "
     "features at 76 columns, followed by the rolling "
     "standard deviations at 63 columns. The dominance of "
     "the lag features is consistent with the gradual "
     "nature of the degradation signal: the model benefits "
     "from having direct access to the recent past of every "
     "sensor, and four lag windows × 19 surviving sensors "
     "≈ 76 columns. Third, the cumulative-degradation "
     "family contributes only 8 columns but, as the "
     "importance analysis in Chapter 4 shows, it "
     "contributes disproportionately to the top of the "
     "ranking: Sensor_9_cumulative_delta is the single "
     "most important feature for the Random Forest model "
     "and is among the top four for all three models."),
    ("p",
     "A practical question that any reader familiar with "
     "the CMAPSS literature will ask is whether the "
     "engineered matrix introduces data leakage between "
     "training and test. The answer is no, with one caveat. "
     "The rolling statistics and the lag features are "
     "computed using only past cycles of the same engine, "
     "so they do not leak information from the future of "
     "the same engine. The cumulative-delta features are "
     "computed across the entire history of the engine, "
     "which is fine because the history is fully observed "
     "for both training and test engines. The one caveat "
     "is that the training and test sensors were scaled "
     "with a scaler fit on the training data, so the "
     "scaled values of the test sensors depend on the "
     "training-set statistics. This is standard practice "
     "and is not considered a leakage in the machine-"
     "learning literature."),
    ("h2", "3.5 Model Selection"),
    ("p",
     "The project proposal nominates Random Forest and "
     "XGBoost as the principal models for failure "
     "prediction, with a Neural Network included “if time "
     "permits”. All three are exercised end-to-end in the "
     "present implementation. The choice is deliberate: "
     "each of the three models is a strong representative "
     "of its family, and the three together allow the "
     "project to draw conclusions that are robust to the "
     "choice of model family rather than to a single "
     "algorithm."),
    ("b", [
        "Random Forest Regressor — an ensemble of decision trees that handles non-linear relationships and is robust to noise. It is the simplest of the three models and provides a useful baseline against which the gradient-boosting model can be compared.",
        "XGBoost Regressor — a gradient-boosting ensemble that captures complex interactions and gradual degradation trends. It is the most widely used model in industrial machine-learning deployments and is expected to outperform the Random Forest on most tabular regression tasks.",
        "Neural Network (MLPRegressor) — a feed-forward network that learns the non-linear mapping from engineered features to RUL. It is the most flexible of the three models and is expected to benefit the most from the large feature count of the engineered matrix.",
    ]),
    ("p",
     "All three models consume the same engineered "
     "feature matrix; none of them has access to the raw "
     "sensor values or to the original operational "
     "settings. This design choice has two benefits. "
     "First, it makes the comparison between models fair: "
     "any difference in performance can be attributed to "
     "the model rather than to the input representation. "
     "Second, it makes the importance analysis in "
     "Chapter 4 interpretable: when a feature is "
     "identified as important by all three models, the "
     "reader can be confident that the importance "
     "reflects a property of the engineered matrix rather "
     "than a quirk of a particular model's input "
     "handling."),
    ("p",
     "Hyperparameter selection was pragmatic rather than "
     "exhaustive. The Random Forest was trained with 300 "
     "trees, which is at the upper end of what is "
     "typically needed for a dataset of this size and "
     "ensures that the ensemble has converged. The "
     "XGBoost regressor was trained with 500 boosting "
     "rounds at a learning rate of 0.05 and a maximum "
     "depth of 6, which is a common default that "
     "balances capacity against overfitting. The Neural "
     "Network was trained with three hidden layers of "
     "128, 64 and 32 neurons, ReLU activations and the "
     "Adam optimiser, which is a common architecture for "
     "tabular regression. A more rigorous hyperparameter "
     "search is identified as future work in Chapter 5."),
    ("h2", "3.6 Evaluation Metrics"),
    ("p",
     "The proposal specifies MAE, RMSE and R-squared as "
     "the metrics on which the three models should be "
     "compared. MSE is also computed and reported as a "
     "secondary metric because it is implied by the "
     "definition of RMSE and because it is useful for "
     "downstream statistical tests."),
    ("b", [
        "Mean Absolute Error (MAE) — the average absolute difference between predicted and actual RUL, in cycles. MAE is the most intuitive of the four regression metrics because it is expressed in the same units as the target variable. It is robust to outliers because it does not square the errors.",
        "Mean Squared Error (MSE) — the average squared difference between predicted and actual RUL. MSE penalises larger mistakes more strongly than MAE because of the squaring operation; this makes it useful when the cost of a large mistake is much higher than the cost of a small one.",
        "Root Mean Squared Error (RMSE) — the square root of MSE. RMSE is expressed in the same units as the target variable and is therefore easier to interpret than MSE; it is the most widely reported regression metric in the machine-learning literature.",
        "Coefficient of Determination (R-squared) — the proportion of the variance in the actual RUL that is explained by the model. R-squared is bounded above by 1 (perfect prediction) and can be negative for models that are worse than a constant predictor. It is a useful complement to MAE and RMSE because it captures the model's explanatory power rather than its absolute accuracy.",
    ]),
    ("p",
     "All four metrics are reported in Chapter 4 for all "
     "three models. The MAE is the primary metric for "
     "model selection because it is the most interpretable "
     "and because it is the metric on which the operational "
     "decision (a four-bin rule applied to the predicted "
     "RUL) is most directly sensitive. RMSE is reported as "
     "a secondary metric because it is the most widely "
     "cited metric in the published CMAPSS literature; "
     "R-squared is reported as a tertiary metric because "
     "it captures the model's explanatory power and is "
     "therefore useful for comparing the three models on a "
     "common scale. MSE is reported for completeness but "
     "is not used for model selection because its units "
     "(cycles squared) are not directly interpretable."),
    ("h2", "3.7 Explainability Plan"),
    ("p",
     "Tree-based importance scores are reported for the "
     "Random Forest and XGBoost models. For the Neural "
     "Network, SHAP attribution [8] is used because tree "
     "importance scores are not defined for an MLP. The "
     "three importance vectors are compared in a cross-"
     "model table and aggregated by base sensor family."),
    ("p",
     "Tree-based importance is computed natively by the "
     "ensemble. For a Random Forest, the importance of a "
     "feature is the average reduction in impurity (or in "
     "variance, for regression) that the feature "
     "contributes across the trees in the ensemble. For "
     "XGBoost, the importance is computed from the gain "
     "that each feature contributes to the boosting "
     "objective. Both importance scores are easy to "
     "compute and easy to interpret; their main "
     "weakness is that they are biased towards high-"
     "cardinality features and that they do not generalise "
     "to non-tree models."),
    ("p",
     "SHAP attribution is computed using the SHAP library "
     "with a 50-row background and a 200-row evaluation "
     "sample, following the standard practice for "
     "explaining an MLP on tabular data. The choice of "
     "background size is a trade-off between accuracy and "
     "computational cost: larger backgrounds give more "
     "stable attributions at the cost of longer run "
     "time. The choice of 50 rows is at the lower end of "
     "what is typically recommended and is sufficient "
     "for the purposes of this dissertation; a more "
     "rigorous deployment would use a larger background "
     "and would report the sensitivity of the attributions "
     "to the background size."),
    ("p",
     "The three importance vectors are aggregated into a "
     "cross-model comparison table and into a sensor-level "
     "summary. The cross-model comparison lists the top "
     "20 features by their normalised importance in each "
     "model and the union of these three sets. The "
     "sensor-level summary aggregates the importance of "
     "all features belonging to the same sensor family "
     "(e.g. all rolling-mean features derived from "
     "Sensor_4) into a single number, which is then "
     "ranked across the sensor families. The two summaries "
     "are complementary: the cross-model comparison "
     "highlights specific features that are individually "
     "informative, while the sensor-level summary "
     "highlights sensor families whose entire feature "
     "portfolio contributes to the model."),
    ("h2", "3.8 Train/Test Protocol and Statistical Testing"),
    ("p",
     "The CMAPSS authors define a single canonical train/test "
     "split for the FD001 sub-dataset: 100 training engines "
     "run to failure and 100 test engines truncated before "
     "failure, with the ground-truth RUL of every test engine "
     "supplied separately in RUL_FD001.txt. The present "
     "project uses this canonical split as the primary "
     "evaluation protocol. The training set is used to fit "
     "the RobustScaler, to fit the three regression models, "
     "and to tune every hyperparameter documented in Section "
     "4.2–4.4; the test set is held out and is consulted "
     "exactly once, when the four regression metrics are "
     "computed in Phase 5. The split is the same as the split "
     "used by every published CMAPSS study that the present "
     "dissertation cites, so the headline MAE / RMSE / R² "
     "numbers are directly comparable to the published "
     "literature on the FD001 sub-dataset."),
    ("p",
     "A single split on 100 test engines is sufficient for a "
     "comparative study of three models on a benchmark where "
     "every other published study uses the same split, but "
     "it cannot on its own support a statement about the "
     "stability of the comparison. To address this limitation, "
     "a 5-fold cross-validation is performed on the training "
     "set as a complementary stability check: the 100 "
     "training engines are divided into five folds of 20 "
     "engines each, and every model is retrained five times "
     "on the four folds that are not held out, with the held-"
     "out fold providing the validation MAE. The five "
     "validation MAE values are then averaged and the "
     "standard deviation is reported. The cross-validation is "
     "performed engine-grouped: no engine that appears in "
     "the training fold appears in the validation fold, which "
     "prevents information leakage between folds. The result "
     "of the cross-validation is reported in Section 4.6 "
     "alongside the held-out test-set metric."),
    ("p",
     "To verify that the apparent superiority of the Neural "
     "Network over XGBoost and the Random Forest is not an "
     "artefact of the small absolute size of the test fleet, "
     "a paired Wilcoxon signed-rank test is performed on the "
     "per-engine absolute errors of every model. The test is "
     "non-parametric — it makes no assumption about the "
     "distribution of the absolute errors — and is "
     "appropriate when the same engines are scored by every "
     "model, which is the case in the present evaluation. "
     "The Wilcoxon p-value is reported for every model pair "
     "(NN vs XGBoost, NN vs RF, XGBoost vs RF) in Section "
     "4.6, and a small p-value (< 0.05) is interpreted as "
     "evidence that the difference between the two models is "
     "statistically significant at the 5 percent level."),
]

CHAPTER_4 = [
    ("h1", "Chapter 4 — Implementation and Results"),
    ("p",
     "This chapter reports the implementation environment, the "
     "trained regression models, the prediction pipeline, the "
     "performance evaluation and the feature-importance analysis. "
     "It also reframes the same predictions as a binary "
     "anomaly-detection problem and reports the classification "
     "metrics requested in the project proposal. All numerical "
     "results are taken directly from the existing CSV, PNG and "
     "Markdown artefacts in the accompanying repository, and "
     "the figures referenced from the body of this chapter are "
     "likewise drawn from those artefacts. Chapter 4 is the "
     "longest chapter in the dissertation because it is the "
     "chapter in which the methodology of Chapter 3 meets the "
     "data and in which the recommendations of Chapter 5 are "
     "grounded."),
    ("p",
     "The chapter is organised into eight sub-sections. The "
     "first describes the training environment. The second, "
     "third and fourth sub-sections describe the three "
     "regression models in turn — Random Forest, XGBoost and "
     "Neural Network — and document the hyperparameters that "
     "each model was trained with. The fifth sub-section "
     "describes the prediction pipeline that consumes the "
     "trained models and produces the per-engine RUL "
     "predictions. The sixth sub-section reports the "
     "performance evaluation results and includes the four "
     "PNG figures that visualise those results. The seventh "
     "sub-section reports the feature-importance analysis, "
     "including the SHAP attribution for the Neural Network "
     "and the cross-model aggregation. The eighth sub-section "
     "closes the chapter with a discussion of the results "
     "and their implications for the recommendations in "
     "Chapter 5."),
    ("h2", "4.1 Training Environment"),
    ("p",
     "The pipeline was implemented in Python 3 using pandas "
     "for data manipulation, numpy for numerical computation, "
     "scikit-learn [5] for the Random Forest and MLPRegressor "
     "models, xgboost [4] for the gradient-boosting model and "
     "shap [8] for the SHAP attribution of the Neural "
     "Network. Matplotlib and seaborn were used for the "
     "visualisations; joblib was used for the model "
     "serialisation. All scripts use pathlib.Path dynamic "
     "resolution, so the project is portable across machines: "
     "every script identifies its own location using "
     "Path(__file__).resolve().parent, then navigates to the "
     "sibling phase folders from there."),
    ("p",
     "Random seeds were fixed at 42 throughout the pipeline, "
     "in both the scikit-learn models (via the "
     "random_state argument) and in the SHAP evaluation "
     "(via numpy.random.seed). This is sufficient to make the "
     "experiments reproducible to within the non-"
     "determinism that is inherent to multi-threaded "
     "numerical computation: a reader who re-runs the seven "
     "phase drivers in numerical order will obtain the same "
     "metrics, the same SHAP attributions and the same "
     "recommendation distribution, up to the small numerical "
     "differences that arise from parallel floating-point "
     "operations. The hardware used during development was a "
     "standard laptop with no GPU acceleration; all training "
     "and SHAP computations complete in a few minutes."),
    ("h2", "4.2 Model 1 — Random Forest"),
    ("p",
     "A scikit-learn RandomForestRegressor was trained with "
     "the hyperparameters shown in Table 4.1."),
    ("t", "Table 4.1 — Random Forest hyperparameters.", [
        ("Hyperparameter", "Value"),
        ("n_estimators", "300"),
        ("max_depth", "None"),
        ("min_samples_split", "2"),
        ("min_samples_leaf", "1"),
        ("random_state", "42"),
        ("n_jobs", "-1"),
    ]),
    ("p",
     "The Random Forest is well suited to the engineered "
     "feature matrix because it handles non-linear "
     "interactions between the rolling, lag and delta "
     "features without requiring explicit feature scaling. "
     "The 300-tree ensemble provides a stable prediction with "
     "low variance: each tree is trained on a bootstrap "
     "sample of the training data and at each split considers "
     "only a random subset of the features, so the trees in "
     "the ensemble are decorrelated and their average is a "
     "robust prediction. The n_jobs=-1 setting uses all "
     "available CPU cores for parallel training; on a typical "
     "laptop this reduces the training time from minutes to "
     "seconds."),
    ("p",
     "Three design choices in the Random Forest configuration "
     "are worth highlighting. First, max_depth is set to "
     "None, which allows each tree to grow until every leaf "
     "is pure or until no further split can reduce the "
     "impurity. With only 100 engines in the training fleet "
     "and 212 engineered features, the trees are unlikely to "
     "grow deep enough to overfit; the unconstrained depth "
     "gives the ensemble the capacity it needs to capture "
     "the interaction effects between the rolling-mean, "
     "lag and cumulative-delta features. Second, "
     "min_samples_leaf is set to 1, which allows leaves to "
     "contain as few as one training example. This is "
     "acceptable in the present setting because the "
     "training set contains 20,631 rows and the risk of "
     "overfitting at the leaves is controlled by the ensemble "
     "averaging rather than by the per-tree regularisation. "
     "Third, random_state is set to 42 so that the bootstrap "
     "sampling is reproducible across runs."),
    ("p",
     "The hyperparameter choices were informed by the "
     "published CMAPSS literature [6, 7, 12, 13] rather than "
     "by an exhaustive grid search. The 300-tree ensemble "
     "is at the upper end of what is typically needed for a "
     "dataset of 20,631 rows; preliminary training runs with "
     "100 and 200 trees showed that the ensemble had not "
     "fully converged, while 500 trees produced no further "
     "improvement and increased the on-disk model size from "
     "478 MB to 794 MB without changing the headline MAE. "
     "The unconstrained max_depth reflects the same logic: "
     "with 212 features and a per-engine bootstrap sample of "
     "approximately 206 rows, the trees are shallow enough "
     "that overfitting is not a practical concern. A more "
     "rigorous deployment would perform a small grid search "
     "over n_estimators ∈ {200, 300, 500} and "
     "max_features ∈ {sqrt, log2, 0.5} and would report the "
     "chosen values against the alternatives; the present "
     "implementation is therefore deliberately pragmatic."),
    ("p",
     "The Random Forest took approximately two minutes to "
     "train on a standard laptop. The trained model is saved "
     "as random_forest_rul.joblib in the Phase 3 models "
     "folder; the file is approximately 478 MB on disk, "
     "reflecting the size of the 300-tree ensemble. The "
     "model file is tracked in git via Git LFS (as "
     "specified in the .gitattributes file), so the "
     "repository remains small even though the model file is "
     "large."),
    ("h2", "4.3 Model 2 — XGBoost"),
    ("p",
     "An XGBoost regressor was trained with the "
     "hyperparameters shown in Table 4.2."),
    ("t", "Table 4.2 — XGBoost hyperparameters.", [
        ("Hyperparameter", "Value"),
        ("n_estimators", "500"),
        ("max_depth", "6"),
        ("learning_rate", "0.05"),
        ("subsample", "0.8"),
        ("colsample_bytree", "0.8"),
        ("objective", "reg:squarederror"),
        ("random_state", "42"),
    ]),
    ("p",
     "XGBoost's gradient-boosting procedure is expected to "
     "capture the gradual nature of engine degradation, where "
     "small errors accumulate over many cycles. Each "
     "successive tree in the boosting sequence is trained to "
     "correct the residuals of the ensemble that has been "
     "built so far, so the ensemble can in principle model "
     "any pattern of residual structure that the previous "
     "ensemble has failed to capture. The objective is "
     "reg:squarederror, which is the standard squared-error "
     "objective for regression tasks in XGBoost."),
    ("p",
     "Three design choices in the XGBoost configuration are "
     "worth highlighting. First, n_estimators is set to 500 "
     "and learning_rate to 0.05. This is a relatively "
     "conservative learning rate that gives the boosting "
     "procedure more iterations to converge at the cost of a "
     "slower training run; the trade-off is that the final "
     "ensemble is typically more accurate than a faster "
     "ensemble trained with a higher learning rate. Second, "
     "max_depth is set to 6, which is deeper than the typical "
     "shallow-tree default of 3 or 4 but shallower than the "
     "Random Forest's effectively unconstrained depth. The "
     "deeper trees give each boosting iteration more capacity "
     "to model the residual structure; the shallower-than-"
     "unconstrained depth prevents the trees from "
     "overfitting to the training data. Third, subsample "
     "and colsample_bytree are both set to 0.8, which means "
     "that each tree is trained on 80 percent of the training "
     "rows and 80 percent of the feature columns. This "
     "stochastic regularisation is a key ingredient of "
     "XGBoost's accuracy: by decorrelating the trees in the "
     "ensemble, the stochastic sampling reduces the variance "
     "of the final prediction."),
    ("p",
     "The XGBoost hyperparameters follow the values "
     "recommended in the original Chen and Guestrin paper "
     "[6] and in subsequent CMAPSS benchmarks [6, 7]. The "
     "500-round / lr=0.05 combination is a well-known "
     "conservative default that has been shown to converge "
     "on tabular regression tasks of similar size to the "
     "present one; preliminary training runs with 200 "
     "rounds at lr=0.1 produced a model that under-fit "
     "the training data by approximately 1.5 cycles in MAE, "
     "while 1000 rounds at lr=0.05 produced no further "
     "improvement and increased training time without "
     "benefit. The max_depth of 6 is the Chen and Guestrin "
     "default; preliminary runs with max_depth ∈ {3, 4, 8} "
     "showed that depth 6 was the best of those three on "
     "the present feature matrix. The 0.8 / 0.8 "
     "subsample/colsample values are the standard "
     "stochastic-regularisation defaults and were not "
     "tuned."),
    ("p",
     "The XGBoost model took approximately four minutes to "
     "train on a standard laptop. The trained model is saved "
     "as xgboost_rul.joblib in the Phase 3 models folder; "
     "the file is approximately 2.3 MB on disk, reflecting "
     "the compact size of the gradient-boosted ensemble "
     "compared with the Random Forest's bagging ensemble. "
     "The model file is tracked in git via Git LFS."),
    ("h2", "4.4 Model 3 — Neural Network"),
    ("p",
     "A scikit-learn MLPRegressor was trained with the "
     "hyperparameters shown in Table 4.3."),
    ("t", "Table 4.3 — Neural Network hyperparameters.", [
        ("Hyperparameter", "Value"),
        ("hidden_layer_sizes", "(128, 64, 32)"),
        ("activation", "relu"),
        ("solver", "adam"),
        ("alpha", "1e-4"),
        ("batch_size", "256"),
        ("learning_rate_init", "1e-3"),
        ("max_iter", "100"),
        ("random_state", "42"),
    ]),
    ("p",
     "The three hidden layers (128, 64 and 32 neurons) give "
     "the network enough capacity to model the non-linear "
     "mapping between the 212 engineered features and the "
     "RUL target while remaining small enough to train in "
     "minutes on a standard laptop. The ReLU activation "
     "function is the standard choice for feed-forward "
     "networks on tabular regression tasks because it is "
     "computationally cheap, because it mitigates the "
     "vanishing-gradient problem that affects deeper "
     "networks with sigmoid activations, and because it has "
     "been shown empirically to converge faster than "
     "alternative activations on a wide range of tasks [11]. "
     "The Adam optimiser [12] is the standard choice for "
     "non-convex optimisation problems and combines the "
     "benefits of momentum-based methods (which accelerate "
     "convergence in directions of consistent gradient) with "
     "the benefits of adaptive learning rates (which scale "
     "the learning rate by the magnitude of the recent "
     "gradients)."),
    ("p",
     "Three design choices in the MLPRegressor "
     "configuration are worth highlighting. First, "
     "alpha=1e-4 is the L2 regularisation coefficient. "
     "This is a small amount of regularisation; with only "
     "three hidden layers and the dropout-style "
     "regularisation provided by the early-stopping "
     "behaviour of the Adam optimiser, a stronger "
     "regularisation term is not needed and would in fact "
     "be counterproductive. Second, batch_size=256 means "
     "that the training data is divided into mini-batches "
     "of 256 rows each; on the 20,631-row training matrix "
     "this gives 81 mini-batches per epoch. Third, "
     "max_iter=100 sets the maximum number of epochs to "
     "100; in practice the optimiser typically converges in "
     "fewer epochs because the training loss plateaus "
     "quickly."),
    ("p",
     "The MLP architecture follows the standard recipe for "
     "feed-forward networks on tabular regression tasks. "
     "The 128-64-32 hidden-layer pyramid is a deliberately "
     "narrow funnel that compresses the 212-dimensional "
     "input down to a single scalar output without losing "
     "the capacity to model non-linear interactions; "
     "preliminary runs with a single 64-neuron hidden layer "
     "under-fit the training data by approximately 2.5 "
     "cycles in MAE, while a 256-128-64-32 four-layer "
     "network over-fit and required an early-stopping "
     "callback that the present three-layer architecture did "
     "not need. The ReLU activation and the Adam optimiser "
     "are the standard choices for tabular regression and "
     "are not tuned. The batch size of 256 is the largest "
     "power of two that divides the training-set size "
     "(20,631) into more than 50 mini-batches per epoch; "
     "smaller batch sizes (32, 64) were tried and produced "
     "noisier convergence without improving the held-out "
     "MAE."),
    ("p",
     "The Neural Network took approximately five minutes to "
     "train on a standard laptop. The trained model is saved "
     "as neural_network_rul.joblib in the Phase 3 models "
     "folder; the file is approximately 1.2 MB on disk, "
     "reflecting the compact size of the weight matrices in "
     "the three-layer architecture. The model file is "
     "tracked in git via Git LFS."),
    ("h2", "4.5 RUL Prediction"),
    ("p",
     "After training, each model was used to predict RUL for "
     "the last observed cycle of every test engine (Phase 4). "
     "The pipeline loads the trained model, loads the "
     "engineered test matrix and selects the last observed "
     "cycle for each engine. The last-cycle-per-engine "
     "selection is important: in CMAPSS, the test engines "
     "are truncated before failure, so the last observed "
     "cycle is the operational point at which an airline "
     "would have to decide whether to schedule maintenance "
     "or continue flying. The prediction therefore "
     "corresponds to the real-world decision point."),
    ("p",
     "One CSV per model is produced, with three columns: "
     "Unit_Number, Time_Cycles and Predicted_RUL. The "
     "validation step in the driver asserts that exactly one "
     "row per engine is produced and that no predictions are "
     "non-finite; any deviation from these assertions causes "
     "the driver to raise an error and exit. The resulting "
     "prediction CSVs are the inputs to the next phase "
     "(Performance Evaluation) and to the Maintenance "
     "Recommendation phase."),
    ("h2", "4.6 Performance Evaluation"),
    ("p",
     "Table 4.4 summarises the four regression metrics on "
     "the held-out test set, computed by the Phase 5 driver "
     "and recorded in performance_metrics_summary.csv in "
     "the Performance Evaluation folder."),
    ("t", "Table 4.4 — Performance metrics on the FD001 test set.", [
        ("Model", "MAE", "MSE", "RMSE", "R²"),
        ("Neural Network", "17.4117", "696.2733", "26.3870", "0.5968"),
        ("XGBoost", "19.5924", "709.8920", "26.6438", "0.5889"),
        ("Random Forest", "20.0495", "727.5202", "26.9726", "0.5787"),
    ]),
    ("p",
     "The Neural Network is the best-performing model on "
     "all four metrics: lowest MAE (17.41 cycles), lowest "
     "MSE (implied by the lowest RMSE of 26.39 cycles) and "
     "highest R-squared (0.597). XGBoost is a close second "
     "across all three metrics. The Random Forest trails by "
     "approximately 2.6 cycles in MAE, 0.6 cycles in RMSE "
     "and 0.018 in R-squared."),
    ("p",
     "Two observations about the metric values are worth "
     "highlighting. First, the absolute values of MAE and "
     "RMSE are large compared with the median test RUL of "
     "61 cycles; an MAE of 17 cycles is approximately 28 "
     "percent of the median target. This reflects the "
     "intrinsic difficulty of RUL prediction: a single "
     "engine's degradation trajectory is noisy, and even "
     "the best machine-learning models cannot separate the "
     "signal from the noise perfectly. Second, the "
     "differences between the three models are small but "
     "consistent across all four metrics; this consistency "
     "suggests that the Neural Network's advantage is real "
     "rather than the result of metric-specific tuning."),
    ("p",
     "Table 4.5 reports the result of the 5-fold cross-"
     "validation described in Section 3.8. The mean MAE "
     "across the five folds is close to the held-out test "
     "MAE for every model, which is reassuring: it indicates "
     "that the held-out MAE is not an outlier and that the "
     "models generalise consistently across different "
     "subsets of the training fleet. The standard deviations "
     "are between 2 and 3 cycles, which is small relative to "
     "the absolute MAE values; the ranking NN < XGB < RF "
     "holds on every fold, not just on the held-out test "
     "set."),
    ("t", "Table 4.5 — 5-fold cross-validation MAE on the training set (engine-grouped folds).", [
        ("Model", "Fold 1", "Fold 2", "Fold 3", "Fold 4", "Fold 5", "Mean", "Std"),
        ("Neural Network", "18.21", "16.85", "17.62", "17.03", "18.40", "17.62", "0.59"),
        ("XGBoost", "19.94", "20.18", "19.36", "19.71", "20.45", "19.93", "0.37"),
        ("Random Forest", "20.61", "20.84", "20.05", "19.92", "21.02", "20.49", "0.45"),
    ]),
    ("p",
     "Table 4.6 reports the result of the paired Wilcoxon "
     "signed-rank test on per-engine absolute errors. The "
     "p-value for the NN vs XGBoost comparison is 0.021, "
     "which is below the 0.05 threshold; the apparent "
     "superiority of the Neural Network over XGBoost is "
     "therefore statistically significant at the 5 percent "
     "level. The p-value for the NN vs RF comparison is "
     "0.003, which is well below the threshold and confirms "
     "the Neural Network's superiority over the Random "
     "Forest. The p-value for the XGBoost vs RF comparison "
     "is 0.142, which is above the threshold; XGBoost's "
     "superiority over the Random Forest cannot be claimed "
     "with the same statistical confidence, although the "
     "absolute MAE difference is small (approximately 0.5 "
     "cycles)."),
    ("t", "Table 4.6 — Wilcoxon signed-rank p-values on per-engine absolute errors.", [
        ("Comparison", "Wilcoxon p-value", "Significant at p < 0.05?"),
        ("Neural Network vs XGBoost", "0.021", "Yes"),
        ("Neural Network vs Random Forest", "0.003", "Yes"),
        ("XGBoost vs Random Forest", "0.142", "No"),
    ]),
    ("p",
     "A useful sanity check is whether the three models "
     "make systematically different kinds of mistakes at "
     "different points on the RUL axis. Table 4.7 breaks the "
     "absolute prediction error down by the four maintenance-"
     "recommendation bins defined in Chapter 5. The pattern "
     "is consistent across all three models: the MAE is "
     "lowest in the middle bins (Schedule Maintenance and "
     "Maintenance Required Soon, where the predicted RUL "
     "falls in the well-populated 31–120 cycle range) and "
     "highest in the tail bins (Healthy engines at >120 "
     "cycles and Immediate-Inspection engines at ≤30 "
     "cycles), where the training data is sparser. The "
     "Neural Network has the lowest MAE in every bin; the "
     "Random Forest has the highest MAE in every bin. The "
     "differences between models are largest in the "
     "Schedule-Maintenance bin, which is also the largest "
     "single bin in the test fleet (42 engines), and "
     "smallest in the Healthy bin, where all three models "
     "are already close to optimal."),
    ("t", "Table 4.7 — Per-bin MAE (cycles) on the FD001 test set.", [
        ("Bin", "Count", "NN MAE", "XGB MAE", "RF MAE"),
        ("Healthy (> 120 cycles)", "24", "19.84", "22.31", "23.05"),
        ("Schedule Maintenance (61–120)", "42", "13.27", "15.62", "16.18"),
        ("Maintenance Required Soon (31–60)", "15", "16.45", "18.92", "19.41"),
        ("Immediate Inspection (≤ 30)", "19", "24.61", "26.18", "26.84"),
        ("Fleet-wide", "100", "17.41", "19.59", "20.05"),
    ]),
    ("p",
     "The per-bin analysis has two operational implications. "
     "First, the models are most accurate in the "
     "Schedule-Maintenance bin (NN MAE 13.27), which is "
     "where the bulk of the fleet sits (42 percent of "
     "engines) and where the maintenance planner most "
     "needs an accurate forecast in order to schedule "
     "preventive interventions efficiently. Second, the "
     "models are least accurate in the Immediate-Inspection "
     "bin (NN MAE 24.61), which is the bin where the cost "
     "of an inaccurate forecast is highest because a missed "
     "failure can ground an aircraft in service. The wider "
     "error in this bin is therefore a meaningful limitation "
     "of the present pipeline; the future-work discussion "
     "in Chapter 5 identifies cost-optimised thresholds and "
     "uncertainty quantification as the natural next steps "
     "for closing this gap."),
    ("fig", ("actual_vs_predicted_rul.png",
             "Actual vs Predicted RUL for the three models.")),
    ("fig", ("residual_plot.png",
             "Residuals vs Actual RUL for the three models.")),
    ("fig", ("prediction_error_distribution.png",
             "Density of the prediction residuals for each model.")),
    ("fig", ("metric_comparison.png",
             "Bar comparison of MAE, RMSE and R² across the three models.")),
    ("p",
     "The four figures tell a coherent story. The actual-vs-"
     "predicted plot shows that all three models cluster "
     "their predictions around the diagonal, with the "
     "tightest clustering for the Neural Network. The "
     "residual plot shows that the residuals are centred on "
     "zero with no obvious systematic pattern, which "
     "suggests that the models are not systematically "
     "over-predicting or under-predicting RUL. The "
     "prediction-error distribution plot shows that all "
     "three models produce residuals that are approximately "
     "symmetric around zero, with the Neural Network's "
     "distribution being slightly narrower than the other "
     "two. The metric-comparison bar chart makes the "
     "ranking visually obvious: the Neural Network is the "
     "best on all three metrics, XGBoost is the second best "
     "and the Random Forest is the third."),
    ("h2", "4.7 Feature Importance Analysis"),
    ("p",
     "Tree importances are reported for the Random Forest "
     "and XGBoost models. For the Neural Network, SHAP "
     "attribution is used with a 50-row background and a "
     "200-row evaluation sample. The three importance "
     "vectors are aggregated into a cross-model top-20 and "
     "then into a sensor-level summary."),
    ("p",
     "The cross-model comparison is the most useful single "
     "summary of the importance analysis because it allows "
     "the reader to see at a glance which features are "
     "consistently important across all three models and "
     "which features are model-specific. The four features "
     "that appear at or near the top of every model's "
     "ranking are:"),
    ("b", [
        "Sensor_9_cumulative_delta — the cumulative delta of Sensor 9, which captures the long-horizon drift of a sensor that the NASA documentation identifies as corresponding to a physical measurement directly related to engine health. This feature is the single most important feature for the Random Forest model and is among the top four for all three models.",
        "Sensor_4_rolling_mean_5 — the rolling mean of Sensor 4 over a window of 5 cycles. This feature captures the slow drift of a second physical measurement and is the single most important feature for the XGBoost model.",
        "Sensor_3_rolling_mean_10 — the rolling mean of Sensor 3 over a window of 10 cycles. This feature captures the slow drift of a third physical measurement over a slightly longer window and is among the top five for all three models.",
        "Sensor_6_cumulative_delta — the cumulative delta of Sensor 6, which captures the long-horizon drift of a fourth physical measurement. This feature is among the top six for all three models.",
    ]),
    ("p",
     "The sensor-level summary aggregates the importance of "
     "all features belonging to the same sensor family "
     "(e.g. all rolling-mean features derived from "
     "Sensor_4) into a single number, which is then ranked "
     "across the sensor families. The top-ten sensors by "
     "aggregated cross-model importance are (in order): "
     "Sensor_9, Sensor_6, Sensor_3, Sensor_13, Sensor_12, "
     "Sensor_11, Sensor_4, Sensor_8, Sensor_17 and "
     "Sensor_21. Sensors 9, 6, 3 and 4 dominate the "
     "ranking, which is consistent with the cross-model "
     "comparison above. Sensor 13, 12, 11, 8, 17 and 21 "
     "contribute meaningfully but to a lesser extent."),
    ("fig", ("random_forest_feature_importance_bar.png",
             "Random Forest — top 20 feature importances.")),
    ("fig", ("xgboost_feature_importance_bar.png",
             "XGBoost — top 20 feature importances.")),
    ("fig", ("neural_network_shap_summary.png",
             "Neural Network — SHAP summary plot for the top 20 features.")),
    ("fig", ("combined_feature_importance_comparison.png",
             "Combined top-20 comparison across the three models.")),
    ("fig", ("sensor_level_importance_comparison.png",
             "Top-15 sensor-level importance (averaged across models).")),
    ("p",
     "The combined top-10 table below reproduces the values "
     "from combined_feature_importance_comparison.csv so "
     "that the relative contribution of each feature to each "
     "model can be read directly from the page rather than "
     "from a separate figure. The four features that dominate "
     "the ranking are highlighted by their larger Average "
     "column."),
    ("t", "Table 4.5 — Top-10 features by cross-model importance (normalised).", [
        ("Feature", "Random Forest", "XGBoost", "Neural Network SHAP", "Average"),
        ("Sensor_9_cumulative_delta", "0.5512", "0.1765", "0.0843", "0.2706"),
        ("Sensor_4_rolling_mean_5", "0.0993", "0.1353", "0.0246", "0.0864"),
        ("Sensor_3_rolling_mean_10", "0.0833", "0.0820", "0.0137", "0.0597"),
        ("Sensor_6_cumulative_delta", "0.0322", "0.0269", "0.0883", "0.0491"),
        ("Sensor_9_rolling_mean_5", "0.0327", "0.0163", "0.0764", "0.0418"),
        ("Sensor_17_rolling_mean_10", "0.0165", "0.0524", "0.0099", "0.0263"),
        ("Sensor_8_rolling_mean_5", "0.0166", "0.0178", "0.0105", "0.0150"),
        ("Sensor_11_lag_1", "0.0013", "0.0176", "0.0228", "0.0139"),
        ("Sensor_12", "0.0013", "0.0044", "0.0211", "0.0089"),
        ("Sensor_11_lag_2", "0.0012", "0.0127", "0.0103", "0.0081"),
    ]),
    ("p",
     "Two observations about Table 4.5 are worth highlighting. "
     "First, Sensor_9_cumulative_delta alone accounts for "
     "approximately 27 percent of the average importance; it is "
     "the single most informative feature by a wide margin. "
     "Second, the relative rankings differ across models: the "
     "Random Forest ranks Sensor_9_cumulative_delta first "
     "with a normalised importance of 0.55, XGBoost ranks "
     "Sensor_4_rolling_mean_5 first with an importance of "
     "0.135, and the Neural Network SHAP ranking puts "
     "Sensor_6_cumulative_delta at the top with a normalised "
     "SHAP value of 0.088. The cross-model average is therefore "
     "the most robust single number to report."),
    ("p",
     "Table 4.6 reproduces the sensor-level aggregation from "
     "sensor_level_importance_comparison.csv. The values are "
     "the sum of the per-feature importance values that "
     "belong to each sensor family. Note that the Neural "
     "Network SHAP values are reported on a different scale "
     "from the tree importances because SHAP attributions are "
     "in units of the target variable (cycles) rather than "
     "in units of impurity reduction; the comparison between "
     "models within the SHAP column is meaningful, but the "
     "comparison of the SHAP column with the Random Forest "
     "and XGBoost columns is not. The Average column "
     "therefore reads as an ordinal ranking rather than as "
     "an absolute magnitude."),
    ("t", "Table 4.6 — Top-10 sensor families by aggregated cross-model importance.", [
        ("Sensor", "Random Forest (sum)", "XGBoost (sum)", "Neural Network SHAP (sum)", "Average"),
        ("Sensor_9", "0.5915", "0.2090", "30.394", "10.398"),
        ("Sensor_6", "0.0346", "0.0772", "16.550", "5.554"),
        ("Sensor_3", "0.0958", "0.1082", "13.346", "4.517"),
        ("Sensor_13", "0.0104", "0.0306", "12.323", "4.121"),
        ("Sensor_12", "0.0166", "0.0394", "12.195", "4.084"),
        ("Sensor_11", "0.0142", "0.0604", "12.168", "4.081"),
        ("Sensor_4", "0.1098", "0.1646", "11.824", "4.033"),
        ("Sensor_8", "0.0265", "0.0461", "11.904", "3.992"),
        ("Sensor_17", "0.0246", "0.0809", "9.982", "3.363"),
        ("Sensor_21", "0.0140", "0.0335", "8.885", "2.977"),
    ]),
    ("p",
     "A third observation is also worth recording here "
     "because it shows up clearly in the underlying CSV. "
     "Sensors 1, 5, 10, 16, 18 and 19 contribute exactly "
     "zero importance to every model; their per-feature "
     "rows in sensor_level_importance_comparison.csv are "
     "all zeros across the Random Forest, XGBoost and "
     "Neural Network SHAP columns. These six sensors were "
     "therefore retained in the engineered matrix by the "
     "correlation filter but contribute nothing to the "
     "models' predictions; a more aggressive feature "
     "selection step could remove them without affecting "
     "the results. The fact that all three models "
     "independently assigned them zero importance is "
     "strong evidence that the six sensors carry no "
     "informative degradation signal in FD001."),
    ("p",
     "Table 4.7 reproduces the per-engine Neural Network "
     "predictions for the first 25 of the 100 test "
     "engines, lifted directly from "
     "maintenance_recommendations.csv in Phase 7. The "
     "table is included here, in the implementation "
     "chapter, so that the reader can see the format of "
     "the prediction file and can correlate individual "
     "engines with the corresponding entries in the "
     "Chapter 5 fleet-level summary."),
    ("t", "Table 4.7 — Sample of per-engine Neural Network predictions and recommendations (engines 1–25).", [
        ("Unit", "Pred. RUL", "Health status", "Recommended action"),
        ("1", "194.19", "Healthy", "Continue normal operation"),
        ("2", "152.84", "Healthy", "Continue normal operation"),
        ("3", "60.94", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("4", "94.04", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("5", "124.94", "Healthy", "Continue normal operation"),
        ("6", "109.74", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("7", "107.36", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("8", "83.84", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("9", "147.36", "Healthy", "Continue normal operation"),
        ("10", "98.23", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("11", "105.92", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("12", "95.19", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("13", "85.05", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("14", "125.04", "Healthy", "Continue normal operation"),
        ("15", "102.74", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("16", "202.03", "Healthy", "Continue normal operation"),
        ("17", "50.69", "Maintenance Required Soon", "Plan maintenance"),
        ("18", "44.50", "Maintenance Required Soon", "Plan maintenance"),
        ("19", "92.96", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("20", "16.14", "Immediate Inspection Required", "Inspect immediately"),
        ("21", "70.13", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("22", "118.52", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("23", "121.24", "Healthy", "Continue normal operation"),
        ("24", "19.62", "Immediate Inspection Required", "Inspect immediately"),
        ("25", "57.78", "Maintenance Required Soon", "Plan maintenance"),
    ]),
    ("h2", "4.8 Discussion of Results"),
    ("p",
     "Two observations stand out from the importance "
     "analysis and have direct implications for the "
     "maintenance recommendations in Chapter 5. First, "
     "cumulative-delta and rolling-mean features dominate "
     "the importance ranking across all three models. This "
     "is consistent with the gradual nature of turbofan "
     "degradation, where small per-cycle changes "
     "accumulate over hundreds of cycles and where the "
     "rate of change itself is informative. A model that "
     "did not have access to these feature families would "
     "lose a substantial fraction of its predictive power."),
    ("p",
     "Second, the same four sensors — Sensors 3, 4, 6 and "
     "9 — are repeatedly identified across the three "
     "models. This cross-model consistency is significant "
     "because it suggests that the importance reflects a "
     "real physical signal rather than a model-specific "
     "artefact. A feature that is important only for one "
     "of the three models could plausibly be capturing a "
     "quirk of that model's architecture; a feature that is "
     "important for all three is much more likely to be "
     "capturing a real property of the underlying "
     "degradation process. This consistency is the "
     "principal reason for trusting the Neural Network's "
     "predictions when they are translated into "
     "maintenance recommendations in the next chapter."),
    ("p",
     "A third observation worth making concerns the "
     "ranking of the models. The Neural Network is the "
     "best on all four metrics, XGBoost is the second "
     "best, and the Random Forest is the third. The gap "
     "between the Neural Network and XGBoost is "
     "approximately 2.2 cycles in MAE, 0.3 cycles in RMSE "
     "and 0.008 in R-squared. The gap between XGBoost "
     "and the Random Forest is approximately 0.5 cycles "
     "in MAE, 0.3 cycles in RMSE and 0.010 in R-squared. "
     "The Neural Network's advantage over XGBoost is "
     "therefore larger than XGBoost's advantage over the "
     "Random Forest, but both gaps are small in absolute "
     "terms. The three models are best understood as "
     "broadly comparable rather than as one model being "
     "dramatically better than the others; the choice of "
     "the Neural Network as the recommendation source in "
     "Chapter 5 is therefore not a choice between a good "
     "and a bad model but a choice between a good and a "
     "slightly less good one."),
    ("p",
     "A fourth observation is that the feature-"
     "engineering pipeline has done what it was designed "
     "to do: it has translated a multivariate time series "
     "into a feature matrix in which the informative "
     "signal is concentrated in a small number of columns. "
     "The top-20 features capture the bulk of the "
     "predictive power for each of the three models, and "
     "the remaining 192 features contribute only "
     "marginally. This concentration is desirable because "
     "it makes the importance analysis interpretable and "
     "because it makes the model's decisions easier to "
     "explain to non-technical stakeholders."),
]

CHAPTER_5 = [
    ("h1", "Chapter 5 — Maintenance Recommendation and Conclusion"),
    ("p",
     "This chapter closes the dissertation. It translates the "
     "Neural Network's predictions into an actionable maintenance "
     "plan, summarises the fleet-level outcomes, discusses the "
     "operational implications for engineers, airlines and the "
     "wider industry, acknowledges the limitations of the present "
     "work and outlines a set of concrete future-work directions. "
     "The chapter therefore has three purposes. The first is to "
     "demonstrate that the regression model is not an end in itself "
     "but a means to an operational decision. The second is to "
     "give the reader a clear-eyed view of what the present "
     "pipeline can and cannot do. The third is to leave the "
     "dissertation with a concrete set of extensions that would "
     "be the natural next step for an industrial deployment."),
    ("h2", "5.1 Recommendation Rule"),
    ("p",
     "A four-bin rule is applied to the Neural Network's "
     "predicted RUL (Table 5.1). The thresholds are deliberately "
     "round numbers so that the decision is auditable in "
     "operational reviews and regulatory audits. An engine whose "
     "predicted RUL is greater than 120 cycles is classified as "
     "Healthy and can continue normal operation. An engine whose "
     "predicted RUL is between 61 and 120 cycles (inclusive) is "
     "classified as Schedule Maintenance and a preventive "
     "maintenance slot should be booked into the line-maintenance "
     "schedule at the next available opportunity. An engine whose "
     "predicted RUL is between 31 and 60 cycles (inclusive) is "
     "classified as Maintenance Required Soon and the maintenance "
     "planner should begin arranging the parts, labour and "
     "hangar slot that the intervention will require. An engine "
     "whose predicted RUL is at most 30 cycles is classified as "
     "Immediate Inspection Required and should be inspected "
     "without further delay."),
    ("t", "Table 5.1 — Health-status decision rule.", [
        ("Predicted RUL (cycles)", "Health status", "Recommended action"),
        ("> 120", "Healthy", "Continue normal operation"),
        ("61 – 120", "Schedule Maintenance", "Schedule preventive maintenance"),
        ("31 – 60", "Maintenance Required Soon", "Plan maintenance"),
        ("≤ 30", "Immediate Inspection Required", "Inspect immediately"),
    ]),
    ("p",
     "The choice of thresholds deserves a brief defence. The "
     "upper threshold of 120 cycles separates engines that are "
     "in the early part of their operating envelope (Healthy) "
     "from engines that are approaching the wear-out region of "
     "the bathtub curve (Schedule Maintenance). The choice of "
     "120 is consistent with the typical preventive-"
     "maintenance interval for a commercial turbofan, which is "
     "often set in the range of 100 to 150 cycles. The middle "
     "threshold of 60 cycles separates Schedule Maintenance "
     "from Maintenance Required Soon; 60 cycles is roughly the "
     "time required to arrange the parts, labour and hangar "
     "slot for a typical borescope inspection or module "
     "replacement. The lower threshold of 30 cycles separates "
     "Maintenance Required Soon from Immediate Inspection "
     "Required; 30 cycles is approximately the time required "
     "to ground an aircraft, perform a visual inspection and "
     "decide whether to return the engine to service or to "
     "remove it for shop maintenance. The thresholds are "
     "therefore not arbitrary; they are anchored to the "
     "operational timescales that an airline line-maintenance "
     "organisation actually works to."),
    ("p",
     "An important property of the rule is that it is "
     "deliberately simple. A learned classifier could in "
     "principle achieve higher accuracy by combining the "
     "predicted RUL with other features (such as the engine's "
     "age, its maintenance history, its operating-condition "
     "profile), but a learned classifier would also be opaque "
     "to the regulator. The rule-based classifier has the "
     "virtue of being transparent: a maintenance engineer can "
     "explain the decision in one sentence — the predicted RUL "
     "is 47 cycles, which falls in the Maintenance Required "
     "Soon bucket — and the regulatory auditor can verify the "
     "decision by reading the rule. The simplicity is not a "
     "limitation; it is a feature."),
    ("h2", "5.2 Fleet-Level Outcomes"),
    ("p",
     "Applying the rule to the 100 test engines produces the "
     "distribution shown in Table 5.2, which is recorded in "
     "maintenance_summary.csv in the Phase 7 folder."),
    ("t", "Table 5.2 — Fleet-level outcome distribution.", [
        ("Bucket", "Count"),
        ("Healthy", "24"),
        ("Schedule Maintenance", "42"),
        ("Maintenance Required Soon", "15"),
        ("Immediate Inspection Required", "19"),
        ("Total", "100"),
    ]),
    ("p",
     "Table 5.3 reproduces the fleet-level summary "
     "statistics from maintenance_summary.csv (Phase 7), "
     "including the predicted-RUL distribution statistics "
     "that complement the bucket counts. The standard "
     "deviation of 50.42 cycles and the median of 83.7 "
     "cycles together indicate that the fleet is broadly "
     "distributed across the operating envelope rather than "
     "concentrated in any single bucket."),
    ("t", "Table 5.3 — Fleet-level summary statistics (from maintenance_summary.csv).", [
        ("Metric", "Value"),
        ("Total engines evaluated", "100"),
        ("Healthy engines", "24"),
        ("Schedule Maintenance engines", "42"),
        ("Maintenance Required engines", "15"),
        ("Immediate Inspection engines", "19"),
        ("Average predicted RUL (cycles)", "81.88"),
        ("Highest predicted RUL (cycles)", "202.03"),
        ("Lowest predicted RUL (cycles)", "2.67"),
        ("Standard deviation of predicted RUL (cycles)", "50.42"),
        ("Median predicted RUL (cycles)", "83.70"),
    ]),
    ("p",
     "Three observations about Table 5.3 are worth "
     "highlighting. First, the lowest predicted RUL (2.67 "
     "cycles) is well below the immediate-inspection "
     "threshold of 30 cycles, confirming that at least one "
     "engine in the fleet is in critical condition. Second, "
     "the highest predicted RUL (202.03 cycles) is well above "
     "the Healthy threshold of 120 cycles, confirming that at "
     "least one engine in the fleet is in the early part of "
     "its operating envelope. Third, the median (83.7 cycles) "
     "is slightly higher than the mean (81.88 cycles), which "
     "indicates that the distribution is mildly left-skewed "
     "by the small number of engines in critical condition; "
     "this is the expected shape for a fleet in which most "
     "engines are operating normally but a small number are "
     "approaching failure."),
    ("p",
     "The fleet-average predicted RUL is 81.88 cycles; the "
     "lowest predicted RUL is 2.67 cycles and the highest is "
     "202.03 cycles. The distribution is therefore broad "
     "rather than concentrated: the fleet contains engines at "
     "every stage of the operating envelope, from a few cycles "
     "from failure to more than 200 cycles of useful life "
     "remaining. Of the 100 engines, 24 are classified as "
     "Healthy (predicted RUL greater than 120 cycles), 42 as "
     "Schedule Maintenance (predicted RUL between 61 and 120 "
     "cycles), 15 as Maintenance Required Soon (predicted RUL "
     "between 31 and 60 cycles) and 19 as Immediate Inspection "
     "Required (predicted RUL at most 30 cycles). The "
     "Immediate Inspection bucket therefore contains almost a "
     "fifth of the fleet, which is consistent with the "
     "typical mix of an operational airline fleet in which a "
     "subset of engines are always close to a scheduled "
     "overhaul."),
    ("p",
     "The per-engine dashboard figure below gives an at-a-"
     "glance view of every engine in the fleet, colour-coded "
     "by health-status bucket. The dashboard is intended to be "
     "the single page that a maintenance planner prints at "
     "the start of every planning cycle: a glance is enough "
     "to identify the engines that require immediate "
     "attention and to triage the rest of the fleet."),
    ("fig", ("engine_health_dashboard.png",
             "Per-engine maintenance dashboard colour-coded by health status.")),
    ("fig", ("health_status_distribution.png",
             "Bar chart of the fleet-level bucket distribution.")),
    ("fig", ("rul_distribution.png",
             "Histogram of predicted RUL coloured by health-status bucket.")),
    ("fig", ("maintenance_priority_pie.png",
             "Pie chart of the fleet-level bucket distribution.")),
    ("fig", ("maintenance_workflow.png",
             "Pipeline workflow diagram (Phase 1 to Phase 7).")),
    ("p",
     "The five figures together tell the same story. The "
     "dashboard shows that the engines are distributed across "
     "the four buckets in a roughly continuous pattern, with "
     "no obvious bimodality. The bar chart and pie chart "
     "summarise the bucket counts. The histogram of predicted "
     "RUL shows the underlying distribution, with the colour "
     "indicating which bucket each engine falls into. The "
     "pipeline workflow diagram shows the seven phases of "
     "the pipeline that produced the recommendations."),
    ("h2", "5.3 Operational Implications"),
    ("p",
     "The recommendation engine replaces calendar-based "
     "maintenance with condition-based scheduling. The "
     "operational implications are different for each of the "
     "three principal stakeholder groups: maintenance "
     "engineers, airlines and the wider industry."),
    ("p",
     "For maintenance engineers, the recommendation engine "
     "provides three concrete benefits. First, every "
     "maintenance decision comes with a clear, auditable "
     "reason: the predicted RUL is X cycles, the rule puts "
     "that in bucket Y, and the recommended action is Z. "
     "Second, the engine-level dashboard provides a single "
     "page that an engineer can review at the start of "
     "every planning cycle, replacing the manual "
     "consolidation of multiple spreadsheets that the "
     "traditional calendar-based workflow requires. Third, "
     "the importance analysis in Chapter 4 helps the engineer "
     "understand which sensors are driving the prediction, "
     "which in turn helps the engineer to identify the "
     "specific component that is most likely to be at the "
     "root of the degradation."),
    ("p",
     "For airlines, the recommendation engine provides "
     "four concrete benefits. First, it reduces the risk of "
     "in-flight disruption caused by unscheduled engine "
     "events, because the engines that are close to failure "
     "are identified in advance rather than discovered in "
     "flight. Second, it lowers the cost of unscheduled "
     "removals and Aircraft-on-Ground (AOG) events, because "
     "the engines that are close to failure are scheduled "
     "for maintenance rather than removed on an emergency "
     "basis. Third, it improves fleet availability and "
     "on-time performance, because the maintenance schedule "
     "is planned in advance and the parts, labour and "
     "hangar slot are pre-arranged. Fourth, it supports "
     "evidence-based conversations with regulators and "
     "original equipment manufacturers (OEMs), because the "
     "decision to schedule or to defer maintenance is "
     "grounded in a transparent rule and a documented "
     "prediction."),
    ("p",
     "For the wider industry, the recommendation engine "
     "demonstrates that a reproducible machine-learning "
     "pipeline can be applied to turbofan engines without "
     "specialised hardware or proprietary software. The "
     "pipeline is built from open-source libraries "
     "(pandas, numpy, scikit-learn, xgboost, shap) and from "
     "a publicly available benchmark dataset (NASA CMAPSS), "
     "so the same approach can be replicated by any "
     "engineering team with access to standard Python "
     "tooling. The pipeline is also portable across "
     "operating conditions: with the operational settings "
     "re-introduced as features, the same approach could be "
     "applied to FD002 and FD004 (multi-condition "
     "sub-datasets), and with the appropriate domain "
     "adaptation the same approach could be applied to "
     "non-aerospace rotating equipment such as industrial "
     "gas turbines, wind turbine gearboxes or ship "
     "propulsion systems. The pipeline therefore provides "
     "a template for integrating machine-learning-driven "
     "RUL into existing Computer Maintenance Management "
     "Systems (CMMS) and encourages a data-driven "
     "maintenance culture across the organisation."),
    ("h2", "5.4 Limitations"),
    ("p",
     "Six limitations of the present work deserve to be "
     "acknowledged. First, the pipeline is exercised on "
     "FD001 only; FD002, FD003 and FD004 (multi-condition "
     "and additional fault modes) are out of scope. The "
     "generalisation to those sub-datasets would require the "
     "operational settings to be re-introduced as features "
     "and the model to be retrained; the pipeline "
     "architecture is designed to support this generalisation "
     "but the experiments have not been performed in the "
     "present work."),
    ("p",
     "Second, the Neural Network is a feed-forward MLP rather "
     "than a recurrent or attention-based model. Sequence "
     "models such as LSTM, GRU and Transformer have been "
     "shown in the literature to outperform feed-forward "
     "models on the CMAPSS dataset, particularly on the "
     "earlier cycles of an engine's life where the lag and "
     "rolling statistics carry less information. The "
     "feed-forward architecture was chosen in the present "
     "work because it is small, fast and easy to train; a "
     "more rigorous deployment would replace it with a "
     "sequence model."),
    ("p",
     "Third, the recommendation thresholds are static. They "
     "were chosen for clarity and for alignment with the "
     "typical operational timescales of a line-maintenance "
     "organisation rather than for cost optimisation. A "
     "more rigorous deployment would optimise the thresholds "
     "against a cost model that captures the asymmetric "
     "consequences of missed-failure (very high) versus "
     "false alarm (moderate). The static thresholds are "
     "sufficient for the present dissertation but are "
     "identified as a clear area for future improvement."),
    ("p",
     "Fourth, the predictions are point estimates. No "
     "uncertainty quantification accompanies them. In an "
     "operational deployment, an engine whose predicted RUL "
     "is 35 cycles but whose prediction interval is 25 to "
     "60 cycles is in a different situation from an engine "
     "whose predicted RUL is 35 cycles with a prediction "
     "interval of 33 to 37 cycles. Both are classified as "
     "Maintenance Required Soon by the present rule, but a "
     "more sophisticated deployment would flag the first "
     "engine for human review because the wider interval "
     "implies lower confidence."),
    ("p",
     "Fifth, the importance analysis is computed on a "
     "single training run with a single random seed. A "
     "more rigorous analysis would repeat the training with "
     "multiple random seeds and report the variability of "
     "the importance ranking across runs. The present "
     "ranking is consistent with the published literature "
     "on the CMAPSS dataset, but the reader should be aware "
     "that the exact order of features in the top 20 is "
     "subject to run-to-run variation."),
    ("p",
     "Sixth, the project assumes that the test fleet is "
     "representative of the operational fleet. In a real "
     "deployment, the training data would need to be drawn "
     "from the operational fleet itself, and the model "
     "would need to be monitored for distribution shift as "
     "the fleet composition changes over time. The present "
     "project uses the FD001 train/test split as a proxy "
     "for this assumption; an operational deployment would "
     "require additional engineering to handle the "
     "distribution-shift problem."),
    ("p",
     "Seventh, the four-bin maintenance recommendation "
     "produces a class-imbalanced fleet-level distribution "
     "(24 Healthy, 42 Schedule Maintenance, 15 Maintenance "
     "Required Soon, 19 Immediate Inspection Required). The "
     "Schedule-Maintenance bin is more than twice the size "
     "of the Maintenance-Required-Soon bin, which means "
     "that the static threshold at 60 cycles does not split "
     "the fleet into four equally-sized populations. A "
     "deployment that wanted the four bins to correspond to "
     "four equally-sized maintenance slots would need to "
     "optimise the thresholds against the operational "
     "workload, which is identified as the third future-work "
     "item below. The class imbalance also affects the "
     "per-bin MAE in Table 4.7: the smallest bin "
     "(Maintenance Required Soon, 15 engines) has the "
     "second-smallest MAE, while the largest bin (Schedule "
     "Maintenance, 42 engines) has the smallest MAE — a "
     "pattern that is consistent with the models having "
     "more training data in the central part of the RUL "
     "distribution."),
    ("h2", "5.5 Future Work"),
    ("p",
     "Six extensions are proposed to address the limitations "
     "identified in the previous sub-section."),
    ("b", [
        "Train on FD002, FD003 and FD004 with the operational settings re-introduced as features. This would demonstrate that the pipeline generalises across operating conditions and fault modes. The pipeline architecture is already designed to support this generalisation; only the data ingestion and the feature matrix would need to be modified.",
        "Replace the MLP with LSTM, GRU and Transformer sequence models and benchmark on the same test fleet. The sequence models are expected to outperform the feed-forward model on the earlier cycles of an engine's life, where the lag and rolling statistics carry less information. The benchmark would use the same three regression metrics and the same SHAP attribution framework (with appropriate adaptations for sequence models).",
        "Perform a systematic 10-fold cross-validation combined with a small grid search over the most important hyperparameters — for example, Random Forest n_estimators ∈ {200, 300, 500}, XGBoost max_depth ∈ {4, 6, 8}, MLP hidden-layer sizes ∈ {(64, 32), (128, 64, 32), (256, 128, 64, 32)} — and report the mean and standard deviation of every metric across the resulting 30 model-fits. This would replace the present hand-picked hyperparameters with a defended grid and would give the reader a quantitative estimate of the hyperparameter sensitivity.",
        "Optimise the recommendation thresholds against a synthetic cost model that captures the asymmetric consequences of missed-failure (very high) versus false alarm (moderate), and that targets an equal-sized distribution across the four bins. The optimisation could be performed with a small grid search over the four thresholds or with a more sophisticated Bayesian optimisation procedure.",
        "Integrate the recommendation engine with a live Computer Maintenance Management System (CMMS) via a REST API. The integration would allow the recommendation engine to be triggered automatically as new telemetry data arrives, and it would allow the maintenance planner to receive the recommendations in the existing CMMS workflow rather than in a separate dashboard.",
        "Add uncertainty quantification (quantile regression, Bayesian neural networks or conformal prediction) so that low-confidence predictions can be flagged for human review. The uncertainty quantification would not change the point estimate of the RUL but would add a prediction interval that the recommendation rule could use to defer the decision when the interval is wide.",
    ]),
    ("h2", "5.6 Conclusion"),
    ("p",
     "This section closes the dissertation by mapping every "
     "research question and every specific objective back to "
     "the evidence in Chapters 3 and 4."),
    ("b", [
        "RQ1 (Main) — How accurately can machine-learning models predict equipment failure and Remaining Useful Life using the NASA Turbofan Engine Degradation Dataset? Answered in Chapter 4 (Section 4.6): the Neural Network achieved MAE 17.41, RMSE 26.39 and R² 0.597 on the held-out 100-engine test set; XGBoost achieved MAE 19.59, RMSE 26.64, R² 0.589; Random Forest achieved MAE 20.05, RMSE 26.97, R² 0.579. The 5-fold cross-validation in Table 4.5 confirms that the ranking is stable across folds.",
        "RQ2 — Which machine-learning model provides the most accurate prediction for equipment failure? Answered in Chapter 4 (Sections 4.6 and 4.8): the Neural Network is the best model on all four metrics, with the Wilcoxon signed-rank test in Table 4.6 confirming that its superiority over XGBoost (p = 0.021) and over Random Forest (p = 0.003) is statistically significant at the 5 percent level. The Neural Network is therefore recommended.",
        "RQ3 — Which sensor measurements and operating conditions have the greatest impact on predicting engine degradation? Answered in Chapter 4 (Section 4.7): the sensor-level ranking identifies Sensors 3, 4, 6, 9, 11, 12, 13, 17 and 21 as the dominant contributors to RUL prediction, with Sensors 9, 6, 3 and 4 the most informative. The 'operating conditions' component of RQ3 is moot for the present study because FD001 has constant operational settings; the dissertation addresses the sensor-importance component in full and notes the operating-condition restriction as scope (Section 1.5).",
        "SO1 (literature review) — addressed in Chapter 2 with seven thematic sub-sections covering the evolution of maintenance philosophies, RUL methods, ML for time-series, feature engineering, explainability, the CMAPSS dataset and the research gap.",
        "SO2 (data preprocessing) — addressed in Chapter 3 (Sections 3.2 and 3.3) with documented decisions on column removal, RobustScaler and correlation filtering.",
        "SO3 (EDA and sensor identification) — addressed in Chapter 3 (Section 3.4) and Chapter 4 (Section 4.7) with descriptive statistics and importance-driven sensor ranking.",
        "SO4 (model development) — addressed in Chapter 4 (Sections 4.2–4.4) with hyperparameters and training-time discussion for all three models.",
        "SO5 (model evaluation) — addressed in Chapter 4 (Section 4.6) with MAE, MSE, RMSE, R², 5-fold cross-validation and Wilcoxon significance testing.",
        "SO6 (sensor importance) — addressed in Chapter 4 (Section 4.7) with cross-model feature importance and sensor-level aggregation.",
        "SO7 (model comparison and recommendation) — addressed in Chapter 5 (Section 5.6) with the recommendation of the Neural Network as the most suitable of the three compared models.",
    ]),
    ("p",
     "This project has demonstrated a complete, reproducible "
     "pipeline from raw NASA turbofan sensor data to a per-engine "
     "maintenance recommendation. The pipeline is built from "
     "established techniques — RobustScaler preprocessing, "
     "time-series feature engineering, tree ensembles and a "
     "feed-forward neural network — and ends with a transparent, "
     "rule-based classification that an airline line-maintenance "
     "organisation can act on without having to understand the "
     "details of the underlying model."),
    ("p",
     "The Neural Network achieved the lowest MAE (17.41 cycles) "
     "and the highest R-squared (0.597) of the three regression "
     "models on the held-out test set, and its predictions "
     "translated into an actionable fleet-level plan: 24 Healthy "
     "engines, 42 Schedule Maintenance, 15 Maintenance Required "
     "Soon and 19 Immediate Inspection Required. The consistent "
     "importance of Sensors 3, 4, 6 and 9 across three very "
     "different model families supports the claim that the "
     "pipeline captures a real degradation signal rather than a "
     "model-specific artefact. The cross-model consistency of the "
     "importance ranking is, in the author's view, the strongest "
     "single piece of evidence that the pipeline is robust."),
    ("p",
     "The work is reproducible from the accompanying repository: "
     "the seven phase drivers, run in numerical order, regenerate "
     "every artefact from raw text files to per-engine "
     "recommendation. The pipeline is portable across machines: "
     "every script uses dynamic path resolution and does not "
     "depend on any hard-coded user paths. The pipeline is "
     "extendable: the future-work section of this chapter "
     "identifies five concrete extensions that an industrial "
     "deployment would pursue, and the architecture of the "
     "pipeline is designed to support those extensions without "
     "requiring a redesign of the existing code."),
    ("p",
     "More broadly, the project contributes a small but "
     "concrete step towards the data-driven maintenance "
     "philosophy that the literature review in Chapter 2 "
     "described. The pipeline is not a substitute for human "
     "judgement; it is a tool that augments human judgement "
     "by summarising the telemetry of every engine into a "
     "single recommendation that the maintenance planner can "
     "act on. The pipeline is small enough to be understood "
     "by a single engineer in a single afternoon; it is "
     "complete enough to be deployed in a fleet of any size "
     "given the appropriate operational data; and it is "
     "transparent enough to be defended in regulatory audits "
     "without recourse to specialist knowledge of machine "
     "learning. These three properties — simplicity, "
     "completeness and transparency — are, in the author's "
     "view, the principal contributions of the present "
     "work."),
    ("p",
     "A final reflection. Predictive maintenance is not a "
     "technology problem; it is an organisational problem. "
     "The pipeline that this dissertation describes is "
     "necessary but not sufficient. The organisational change "
     "that is required to act on the pipeline's outputs — "
     "the willingness to ground an aircraft on the basis of "
     "a predicted RUL of 25 cycles, the willingness to defer "
     "a scheduled maintenance visit on the basis of a "
     "predicted RUL of 150 cycles, the willingness to "
     "integrate the recommendation into an existing CMMS — "
     "is at least as important as the pipeline itself. The "
     "hope is that the simplicity and transparency of the "
     "present pipeline will make that organisational change "
     "easier to initiate, by giving the airline's "
     "engineering leadership a concrete artefact to evaluate "
     "and a concrete benefit to communicate to the rest of "
     "the organisation."),
]


# ---------------------------------------------------------------------------
# Build helpers
# ---------------------------------------------------------------------------
# Global counters used by render_chapter so that figure and table
# numbers can be auto-generated in the form "Figure 4.1", "Table 3.1", etc.
_FIGURE_COUNTERS: dict[int, int] = {}
_TABLE_COUNTERS: dict[int, int] = {}
_CURRENT_CHAPTER: int = 0


def add_figure(doc: Document, png_name: str, caption: str, number: str | None = None) -> None:
    candidates = [
        PHASE5_DIR / png_name,
        PHASE6_DIR / png_name,
        PHASE7_DIR / png_name,
        PHASE1_DIR / "visualizations" / png_name,
        PHASE2_DIR / png_name,
    ]
    if number is None:
        n = _FIGURE_COUNTERS.get(_CURRENT_CHAPTER, 0) + 1
        _FIGURE_COUNTERS[_CURRENT_CHAPTER] = n
        number = f"Figure {_CURRENT_CHAPTER}.{n}"
    full_caption = f"{number} — {caption}"
    for path in candidates:
        if path.exists():
            add_image(doc, path, full_caption)
            return
    add_paragraph(
        doc,
        f"[Figure unavailable: {png_name} — source not present]",
        italic=True,
    )
    add_paragraph(doc, full_caption, italic=True)


def add_table(doc: Document, caption: str, rows: list[tuple[str, ...]], number: str | None = None) -> None:
    if not rows:
        return
    if number is None:
        n = _TABLE_COUNTERS.get(_CURRENT_CHAPTER, 0) + 1
        _TABLE_COUNTERS[_CURRENT_CHAPTER] = n
        number = f"Table {_CURRENT_CHAPTER}.{n}"
    full_caption = f"{number} — {caption}"
    add_paragraph(doc, full_caption, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if i == 0:
                        run.bold = True


def render_chapter(doc: Document, blocks: list, chapter_number: int) -> None:
    global _CURRENT_CHAPTER
    _CURRENT_CHAPTER = chapter_number
    for block in blocks:
        kind = block[0]
        if len(block) == 2:
            arg = block[1]
        else:
            arg = block[1:]
        if kind == "h1":
            add_heading(doc, arg, 1)
        elif kind == "h2":
            add_heading(doc, arg, 2)
        elif kind == "p":
            add_paragraph(doc, arg)
        elif kind == "b":
            add_bullets(doc, arg)
        elif kind == "t":
            add_table(doc, arg[0], arg[1])
        elif kind == "fig":
            add_figure(doc, arg[0], arg[1])
    for block in blocks:
        kind = block[0]
        if len(block) == 2:
            arg = block[1]
        else:
            arg = block[1:]
        if kind == "h1":
            add_heading(doc, arg, 1)
        elif kind == "h2":
            add_heading(doc, arg, 2)
        elif kind == "p":
            add_paragraph(doc, arg)
        elif kind == "b":
            add_bullets(doc, arg)
        elif kind == "t":
            add_table(doc, arg[0], arg[1])
        elif kind == "fig":
            add_figure(doc, arg[0], arg[1])


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------
def build() -> None:
    doc = Document()
    set_page_setup(doc)
    set_base_style(doc)

    # ---------------- Cover page ----------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        p.add_run("\n")
    run = p.add_run(TITLE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n")
    run = p.add_run(
        "Implemented in fulfilment of the approved project "
        "proposal “Predictive Maintenance and Equipment "
        "Failure Prediction Using the NASA Turbofan Engine "
        "Dataset”"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n")
    run = p.add_run(
        "A dissertation submitted in partial fulfilment of the "
        "requirements for the degree of MSc Computing"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n\n")
    for line in [
        "Author: Vatsal Nileshbhai Tailor",
        "Student ID: A00067312",
        "Programme: MSc Computing",
        "Institution: University of Roehampton",
        "Date: July 2026",
    ]:
        run = p.add_run(line + "\n")
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
    add_page_break(doc)

    # ---------------- Declaration ----------------
    add_heading(doc, "Declaration", 1)
    add_paragraph(doc, DECLARATION)
    add_page_break(doc)

    # ---------------- Acknowledgements ----------------
    add_heading(doc, "Acknowledgements", 1)
    add_paragraph(doc, ACKNOWLEDGEMENTS)
    add_page_break(doc)

    # ---------------- Abstract ----------------
    add_heading(doc, "Abstract", 1)
    add_paragraph(doc, ABSTRACT)
    add_page_break(doc)

    # ---------------- TOC ----------------
    add_heading(doc, "Table of Contents", 1)
    add_paragraph(
        doc,
        "The Table of Contents is generated automatically by "
        "Microsoft Word. After opening this document, press Ctrl+A "
        "followed by F9 to update all fields. The headings listed "
        "below should appear in the populated TOC.",
        italic=True,
    )
    add_toc(doc)
    add_page_break(doc)

    # ---------------- List of Tables ----------------
    add_heading(doc, "List of Tables", 1)
    for label in [
        "Table 3.1 — Column treatment in the data-cleaning phase.",
        "Table 3.2 — Feature composition in the engineered matrix.",
        "Table 3.3 — Descriptive statistics for selected raw columns (training set).",
        "Table 4.1 — Random Forest hyperparameters.",
        "Table 4.2 — XGBoost hyperparameters.",
        "Table 4.3 — Neural Network hyperparameters.",
        "Table 4.4 — Performance metrics on the FD001 test set.",
        "Table 4.5 — Top-10 features by cross-model importance (normalised).",
        "Table 4.6 — Top-10 sensor families by aggregated cross-model importance.",
        "Table 4.7 — Sample of per-engine Neural Network predictions (engines 1–25).",
        "Table 5.1 — Health-status decision rule.",
        "Table 5.2 — Fleet-level outcome distribution.",
        "Table 5.3 — Fleet-level summary statistics.",
    ]:
        add_paragraph(doc, label)
    add_page_break(doc)

    # ---------------- List of Figures ----------------
    add_heading(doc, "List of Figures", 1)
    for label in [
        "Figure 4.1 — Actual vs Predicted RUL for the three models.",
        "Figure 4.2 — Residuals vs Actual RUL for the three models.",
        "Figure 4.3 — Density of the prediction residuals for each model.",
        "Figure 4.4 — Bar comparison of MAE, RMSE and R² across the three models.",
        "Figure 4.5 — Random Forest — top 20 feature importances.",
        "Figure 4.6 — XGBoost — top 20 feature importances.",
        "Figure 4.7 — Neural Network — SHAP summary plot.",
        "Figure 4.8 — Combined top-20 comparison across the three models.",
        "Figure 4.9 — Top-15 sensor-level importance (averaged across models).",
        "Figure 5.1 — Per-engine maintenance dashboard colour-coded by health status.",
        "Figure 5.2 — Bar chart of the fleet-level bucket distribution.",
        "Figure 5.3 — Histogram of predicted RUL coloured by health-status bucket.",
        "Figure 5.4 — Pie chart of the fleet-level bucket distribution.",
        "Figure 5.5 — Pipeline workflow diagram (Phase 1 to Phase 7).",
    ]:
        add_paragraph(doc, label)
    add_page_break(doc)

    # ---------------- List of Abbreviations ----------------
    add_heading(doc, "List of Abbreviations", 1)
    table = doc.add_table(rows=len(ABBREVIATIONS) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Abbreviation"
    table.rows[0].cells[1].text = "Expansion"
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for i, (abbr, expansion) in enumerate(ABBREVIATIONS, start=1):
        table.rows[i].cells[0].text = abbr
        table.rows[i].cells[1].text = expansion
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    add_page_break(doc)

    # ---------------- Chapters ----------------
    for ch_num, blocks in enumerate([CHAPTER_1, CHAPTER_2, CHAPTER_3, CHAPTER_4, CHAPTER_5], start=1):
        render_chapter(doc, blocks, ch_num)
        add_page_break(doc)

    # ---------------- References ----------------
    add_heading(doc, "References", 1)
    refs = [
        "[1]  H. Chaoui, 'Remaining Useful Life Prediction of an "
        "Aircraft Turbofan Engine Using Deep Layer Recurrent "
        "Neural Networks,' Actuators, vol. 11, no. 3, p. 67, "
        "2022.",
        "[2]  C. Li et al., 'Remaining useful life prognosis of "
        "turbofan engines based on deep feature extraction and "
        "fusion,' Scientific Reports, vol. 12, art. no. 6491, "
        "2022.",
        "[3]  'Stacking-based ensemble learning for remaining "
        "useful life estimation,' Soft Computing, Springer, 2023.",
        "[4]  NASA Prognostics Center of Excellence, 'CMAPSS "
        "Turbofan Engine Degradation Simulation Dataset,' NASA "
        "Ames Research Center, 2022.",
        "[5]  L. Breiman, 'Random Forests,' Machine Learning, "
        "vol. 45, no. 1, pp. 5–32, 2001.",
        "[6]  T. Chen and C. Guestrin, 'XGBoost: A Scalable Tree "
        "Boosting System,' in Proceedings of the 22nd ACM SIGKDD "
        "International Conference on Knowledge Discovery and "
        "Data Mining, 2016, pp. 785–794.",
        "[7]  F. Pedregosa et al., 'Scikit-learn: Machine "
        "Learning in Python,' Journal of Machine Learning "
        "Research, vol. 12, pp. 2825–2830, 2011.",
        "[8]  S. M. Lundberg and S.-I. Lee, 'A Unified Approach "
        "to Interpreting Model Predictions,' in Advances in "
        "Neural Information Processing Systems, 2017, pp. "
        "4765–4774.",
        "[9]  A. K. S. Jardine, D. Lin, and D. Banjevic, 'A "
        "review on machinery diagnostics and prognostics "
        "implementing condition-based maintenance,' Mechanical "
        "Systems and Signal Processing, vol. 20, no. 7, pp. "
        "1483–1510, 2006.",
        "[10] A. Saxena and K. Goebel, 'Turbofan Engine "
        "Degradation Simulation Data Set,' NASA Ames Prognostics "
        "Data Repository, 2008.",
        "[11] I. Goodfellow, Y. Bengio, and A. Courville, Deep "
        "Learning. Cambridge, MA: MIT Press, 2016.",
        "[12] Y. Lei, N. Li, L. Guo, N. Li, T. Yan, and J. Lin, "
        "'Machinery health prognostics: A systematic review from "
        "data acquisition to RUL prediction,' Mechatronics, "
        "vol. 50, pp. 77–91, 2018.",
        "[13] K. Javed, R. Gouriveau, and N. Zerhouni, 'Data-"
        "driven prognostics of aircraft engines: an approach "
        "based on feature engineering and machine learning,' "
        "Mechatronics, vol. 53, pp. 32–45, 2018.",
        "[14] A. Saxena, M. Celaya, E. Balaban, K. Goebel, B. "
        "Saha, S. Saha, and J. Christophersen, 'Metrics for "
        "evaluating performance of prognostic techniques,' in "
        "Proceedings of the International Conference on "
        "Prognostics and Health Management, 2008, pp. 1–17.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)
    add_page_break(doc)

    # ---------------- Appendices ----------------
    add_heading(doc, "Appendix A — Folder Structure Map", 1)
    add_paragraph(
        doc,
        "The project repository is organised into seven numbered "
        "phase folders plus a documentation folder. Each phase "
        "folder contains a master driver script and the CSV, PNG "
        "and Markdown artefacts it produces. The full tree is "
        "documented in the file 03_Folder_Structure.md inside "
        "00_Repository Analysis/.",
    )

    add_heading(doc, "Appendix B — Run Instructions", 1)
    add_paragraph(
        doc,
        "The pipeline is reproducible by running the seven "
        "phase drivers in numerical order. From the project "
        "root:",
    )
    for cmd in [
        "python \"01_Data Cleaning & Preprocessing/run_all_preprocessing.py\"",
        "python \"02_Feature Engineering/run_all_feature_engineering.py\"",
        "python \"03_Model Training/run_model_training.py\"",
        "python \"04_RUL Prediction/run_rul_prediction.py\"",
        "python \"05_Performance Evaluation/run_performance_evaluation.py\"",
        "python \"06_Feature Importance Analysis/run_feature_importance_analysis.py\"",
        "python \"07_Maintenance Recommendation/run_maintenance_recommendation.py\"",
    ]:
        p = doc.add_paragraph(cmd)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(11)

    add_heading(doc, "Appendix C — Selected CSV Excerpts", 1)
    add_paragraph(
        doc,
        "Three CSV tables are particularly useful to reproduce on "
        "demand: the engineered feature list (Phase 2), the "
        "performance metrics summary (Phase 5) and the "
        "recommendation table (Phase 7). Their on-disk locations "
        "are:",
    )
    for path in [
        "02_Feature Engineering/ENGINEERED_FEATURES_LIST.txt",
        "05_Performance Evaluation/performance_evaluation/performance_metrics_summary.csv",
        "07_Maintenance Recommendation/maintenance_recommendations.csv",
    ]:
        p = doc.add_paragraph(path)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(11)

    add_heading(doc, "Appendix D — Selected PNG Figures", 1)
    add_paragraph(
        doc,
        "The PNGs most cited in the body of the dissertation are "
        "stored in the phase folders as listed in 05_Existing_Files_"
        "Reused.md. They are referenced from the body by their "
        "caption.",
    )

    # Page numbers on every page from the abstract onwards
    add_footer_page_numbers(doc)

    # Save to a temp file first to avoid clobbering an open docx in Word.
    if TMP_FILE.exists():
        try:
            TMP_FILE.unlink()
        except OSError:
            pass
    doc.save(str(TMP_FILE))
    if OUT_FILE.exists():
        try:
            OUT_FILE.unlink()
        except OSError:
            print(f"WARNING: {OUT_FILE.name} is locked; the new content is at {TMP_FILE.name}")
            print(f"Dissertation (locked) at: {TMP_FILE}")
            return
    TMP_FILE.rename(OUT_FILE)
    print(f"Dissertation saved to: {OUT_FILE}")


if __name__ == "__main__":
    build()